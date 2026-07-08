import os
import json
from datetime import datetime
from llms import call_llm_api, call_llm_api_stream
from typing import Dict, Any, List
import zipfile
import re
from io import BytesIO
import hashlib

# ==================== 路径配置（通过环境变量覆盖，默认使用相对路径）====================
VECTOR_DB_PATH = os.getenv("VECTOR_DB_PATH", "./DB/vector_db")
CHROMA_DB_PATH = os.getenv("CHROMA_DB_PATH", "./DB/chroma_db")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "./models/Ceceliachenen/paraphrase-multilingual-MiniLM-L12-v2")

# 通义千问Embedding配置
DASHSCOPE_EMBEDDING_MODEL = "text-embedding-v3"  # 通义文本嵌入模型111

# OCR 并发控制（熔断机制）
import asyncio
_OCR_SEMAPHORE = None

def _get_ocr_semaphore():
    global _OCR_SEMAPHORE
    if _OCR_SEMAPHORE is None:
        _OCR_SEMAPHORE = asyncio.Semaphore(int(os.getenv("OCR_MAX_CONCURRENT", "5")))
    return _OCR_SEMAPHORE

OCR_TIMEOUT = int(os.getenv("OCR_TIMEOUT", "60"))  # 单张图片超时（秒）
OCR_MAX_RETRIES = int(os.getenv("OCR_MAX_RETRIES", "2"))  # 失败重试次数

def get_current_datetime():
    now = datetime.now()
    return now.strftime("%Y%m%d%H%M%S")


def _fix_json_control_chars(json_str: str) -> str:
    """
    修复JSON字符串中的控制字符问题
    大模型返回的JSON中，字符串字段可能包含真实的换行符、制表符等，
    需要将其转换为转义序列才能正确解析
    """
    import re

    # 在JSON字符串值内部，将真实的控制字符替换为转义序列
    # 匹配JSON字符串值: "..." 但不匹配已转义的
    def fix_string_value(match):
        s = match.group(0)
        # 替换未转义的控制字符
        # 注意：要保留已经转义的 \n \r \t
        result = []
        i = 0
        while i < len(s):
            if s[i] == '\\' and i + 1 < len(s):
                # 已转义的字符，保留
                result.append(s[i:i+2])
                i += 2
            elif s[i] == '\n':
                result.append('\\n')
                i += 1
            elif s[i] == '\r':
                result.append('\\r')
                i += 1
            elif s[i] == '\t':
                result.append('\\t')
                i += 1
            elif ord(s[i]) < 32:
                # 其他控制字符替换为空格
                result.append(' ')
                i += 1
            else:
                result.append(s[i])
                i += 1
        return ''.join(result)

    # 匹配JSON字符串（包括多行）
    # 这个正则匹配 "..." 形式的字符串，处理转义引号
    pattern = r'"(?:[^"\\]|\\.)*"'

    try:
        fixed = re.sub(pattern, fix_string_value, json_str, flags=re.DOTALL)
        return fixed
    except Exception:
        # 如果正则处理失败，尝试简单替换
        return json_str.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')


# 添加新函数：构建系统提示词
def build_system_message(
        requirement_text: str,
        test_level: str,
        test_module: str,
        test_case_count: int,
        prompt: str = ""
) -> str:
    """
    构建系统提示词（不调用大模型）

    Args:
        requirement_text: 需求描述文本
        test_level: 测试级别
        test_priority: 测试优先级
        test_case_count: 测试用例数量
        prompt: 额外的提示词

    Returns:
        str: 完整的系统提示词
    """
    if prompt and test_case_count > 0:
        # 使用用户提供的提示词模板
        return prompt.format(
            requirement_text=requirement_text,
            test_level=test_level,
            test_module=test_module,
            test_case_count=test_case_count
        )
    elif prompt and test_case_count == -1:
        # 1. 把模板里的“请生成 X 条”替换成让 AI 自己决定
        prompt = prompt.replace(
            "请生成 {test_case_count} 个测试用例",
            "请根据你对需求描述的理解，全面输出合理的测试用例条数"
        )

        # 2. 组装其余变量（不再给 test_case_count 赋值）
        return prompt.format(
            requirement_text=requirement_text,
            test_level=test_level,
            test_module=test_module
            # test_case_count 不再传
        )
    else:
        # 使用默认提示词模板
        return f"""
你是一名资深测试工程师，擅长从需求文档中提取关键功能点并设计覆盖全场景的测试用例。请根据用户提供的需求文档，严格按照以下要求输出测试用例：

## 输入处理要求
1. 逐项解析需求文档中的功能点、业务规则和约束条件
2. 识别显性需求（文档明确描述）和隐性需求（行业常识/用户体验）
3. 特别注意边界条件、异常流程和关联功能影响

需求描述:
{requirement_text}


测试用例数量: 请生成 {test_case_count} 个测试用例


请确保测试用例全面覆盖以下测试类型:
1. 功能测试 - 验证功能是否按照需求正确实现
2. 主流程测试 - 验证核心业务流程正常工作
3. 边界条件测试 - 验证系统在极限值和边界情况下的表现
4. 异常情况测试 - 验证系统对错误输入和异常情况的处理
5. 用户界面测试 - 验证UI元素的正确显示和交互(如适用)

每个测试用例必须包含以下信息:
- 模块名称
- 唯一的测试用例ID (格式为TC-xxx，从001开始递增)
- 测试优先级 (与输入参数保持一致)
- 清晰简洁的测试标题
- 详细的前置条件
- 明确的测试步骤 (每个步骤单独一行，使用换行符分隔)
- 具体的预期结果 (每项单独一行，使用换行符分隔)

请直接输出符合以下格式的JSON，不要包含任何额外的说明、注释或Markdown标记:

{{
  "test_cases": [
    {{
      "test_module": "{test_module}",
      "case_id": "TC-001",
      "priority": "高/中/低",
      "title": "测试用例标题",
      "precondition": "测试前置条件",
      "steps": "1. 打开系统登录页面\n2. 输入用户名和密码\n3. 点击登录按钮",
      "expected_result": "1. 登录页面正常显示\n2. 用户名和密码输入框可用\n3. 成功登录进入首页"
    }}
  ]
}}

重要提示:
1. 确保生成的JSON格式完全有效且可直接解析
2. 所有字段必须填写完整，不能有空值
3. 不要输出JSON Schema或其他格式
4. 测试用例应直接关联需求，确保需求的每个方面都有测试覆盖
5. 测试步骤必须每个步骤单独一行，使用换行符 \\n 分隔，不要用分号或其他符号分隔
6. 预期结果必须每项单独一行，使用换行符 \\n 分隔，不要用分号或其他符号分隔
7. 换行符在JSON中表示为 \\n（两个字符：反斜杠+n）
"""


async def generate_test_cases_stream(
        requirement_text: str,
        test_level: str,
        test_module: str,
        test_case_count: int,
        system_message: str,
        provider: str | None = None,
        prompt: str = ""
):
    """
    流式生成测试用例，逐块返回内容
    
    现在会从流式API最后一条消息中提取 token_usage 并传递到 complete 事件中

    Yields:
        dict: 包含类型和内容的字典
            - {"type": "chunk", "content": str}: 流式文本片段
            - {"type": "complete", "content": str, "test_cases": dict, "_token_usage": dict}: 完成时返回完整内容和解析后的测试用例
            - {"type": "error", "content": str}: 错误信息
    """
    # 如果未提供system_message，则构建它
    if system_message is None:
        system_message = build_system_message(
            requirement_text, test_level, test_module, test_case_count, prompt
        )

    full_content = ""
    token_usage = {}
    try:
        print("发送给到大模型的提示词为: " + system_message)
        # 流式调用大模型API
        async for chunk in call_llm_api_stream(system_message, provider=provider):
            # 检查是否为错误
            if isinstance(chunk, dict) and "error" in chunk:
                yield {"type": "error", "content": chunk.get("error")}
                return

            # 检查是否为 token_usage 信息（流式API最后一条消息）
            if isinstance(chunk, dict) and "_token_usage" in chunk:
                token_usage = chunk["_token_usage"]
                continue

            # 累积内容
            full_content += chunk
            # 返回流式片段
            yield {"type": "chunk", "content": chunk}

        # 流式完成后，解析完整的JSON
        try:
            # 先修复JSON中的控制字符
            fixed_content = _fix_json_control_chars(full_content)
            # 尝试直接解析整个内容
            test_cases = json.loads(fixed_content)

            # 验证结构是否正确
            if "test_cases" not in test_cases or not isinstance(test_cases["test_cases"], list):
                raise ValueError("生成的测试用例格式不正确")

            yield {"type": "complete", "content": full_content, "test_cases": test_cases, "_token_usage": token_usage}
        except json.JSONDecodeError:
            # 如果直接解析失败，尝试提取JSON部分
            json_pattern = r'\{[\s\S]*\}'
            match = re.search(json_pattern, full_content)

            if match:
                json_str = match.group(0)
                # 修复JSON中的控制字符
                fixed_json_str = _fix_json_control_chars(json_str)
                test_cases = json.loads(fixed_json_str)

                # 验证结构是否正确
                if "test_cases" not in test_cases or not isinstance(test_cases["test_cases"], list):
                    raise ValueError("生成的测试用例格式不正确")

                yield {"type": "complete", "content": full_content, "test_cases": test_cases, "_token_usage": token_usage}
            else:
                yield {"type": "error", "content": "无法从API响应中提取JSON"}

    except Exception as e:
        yield {"type": "error", "content": f"生成测试用例失败: {str(e)}"}


async def parse_markdown_with_images(content: str, base_path: str = None, vision_provider: str = "aliyun") -> str:
    """
    解析 Markdown 文件中的图片引用，调用视觉大模型识别图片内容，
    并将识别结果还原到原文本位置，标注为【图片解析】

    Args:
        content: Markdown 文件的文本内容
        base_path: MD 文件所在的基础路径（用于解析相对路径图片）
        vision_provider: 视觉模型提供商，"aliyun" 或 "doubao"

    Returns:
        str: 处理后的文本，图片引用被替换为识别结果
    """
    import aiohttp
    import asyncio
    from urllib.parse import urlparse, urljoin

    # 匹配 Markdown 图片语法: ![alt](path/url) 或 ![alt](path/url "title")
    image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"[^"]*")?\)')

    # 找到所有图片引用
    matches = list(image_pattern.finditer(content))

    if not matches:
        return content

    print(f"[MD图片解析] 找到 {len(matches)} 个图片引用")

    async def fetch_and_ocr_image(match) -> tuple:
        """获取图片并进行 OCR 识别"""
        alt_text = match.group(1)
        image_path = match.group(2)
        original_text = match.group(0)

        try:
            image_bytes = None

            # 判断是 URL 还是本地路径
            parsed = urlparse(image_path)
            if parsed.scheme in ('http', 'https'):
                # 网络图片，使用 aiohttp 下载
                async with aiohttp.ClientSession() as session:
                    async with session.get(image_path, timeout=aiohttp.ClientTimeout(total=30)) as response:
                        if response.status == 200:
                            image_bytes = await response.read()
                        else:
                            print(f"[MD图片解析] 下载图片失败: {image_path}, 状态码: {response.status}")
                            return match.start(), original_text, None
            else:
                # 本地图片路径
                local_path = image_path

                # 如果是相对路径，结合 base_path
                if base_path and not os.path.isabs(image_path):
                    local_path = os.path.join(base_path, image_path)

                # 处理 Windows 路径
                local_path = local_path.replace('/', os.sep).replace('\\', os.sep)

                if os.path.exists(local_path):
                    with open(local_path, 'rb') as f:
                        image_bytes = f.read()
                else:
                    print(f"[MD图片解析] 本地图片不存在: {local_path}")
                    return match.start(), original_text, None

            if image_bytes:
                # 调用 OCR 识别
                ocr_result = await ocr_image_async(image_bytes, vision_provider=vision_provider)

                if ocr_result and ocr_result.strip():
                    # 构建替换文本，保留原图片引用并添加解析结果
                    alt_desc = f"（{alt_text}）" if alt_text else ""
                    replacement = f"\n\n【图片解析开始】{alt_desc}\n{ocr_result}\n【图片解析结束】\n\n"
                    return match.start(), original_text, replacement
                else:
                    print(f"[MD图片解析] 图片 OCR 无结果: {image_path}")
                    return match.start(), original_text, None

        except asyncio.TimeoutError:
            print(f"[MD图片解析] 下载图片超时: {image_path}")
        except Exception as e:
            print(f"[MD图片解析] 处理图片失败: {image_path}, 错误: {str(e)}")

        return match.start(), original_text, None

    # 并发处理所有图片
    tasks = [fetch_and_ocr_image(m) for m in matches]
    results = await asyncio.gather(*tasks)

    # 按位置倒序排列，从后往前替换（避免位置偏移）
    results_sorted = sorted(results, key=lambda x: x[0], reverse=True)

    # 执行替换
    result_content = content
    for start_pos, original_text, replacement in results_sorted:
        if replacement:
            result_content = result_content.replace(original_text, replacement, 1)

    success_count = sum(1 for _, _, r in results if r is not None)
    print(f"[MD图片解析] 成功解析 {success_count}/{len(matches)} 个图片")

    return result_content


async def parse_pdf_to_text(file_bytes: bytes, enable_ocr: bool = True) -> str:
    """
    解析 PDF 文件，提取文本内容和图片（OCR识别）
    图片 OCR 结果会被插入到对应页面的位置，标注为【图片解析】

    Args:
        file_bytes: PDF 文件的字节内容
        enable_ocr: 是否启用图片 OCR 识别

    Returns:
        str: 提取的文本内容（包含 OCR 识别的图片内容，位于对应页面）
    """
    try:
        import fitz  # PyMuPDF
        import asyncio

        # 从字节创建 PDF 文档
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # 按页面组织内容：{page_num: {"text": str, "images": [(img_index, image_bytes)]}}
        page_contents = {}

        for page_num, page in enumerate(doc):
            page_contents[page_num] = {
                "text": "",
                "images": []
            }

            # 提取文本内容
            page_text = page.get_text()
            if page_text.strip():
                page_contents[page_num]["text"] = page_text

            # 提取图片
            if enable_ocr:
                page_images = page.get_images(full=True)
                for img_index, img in enumerate(page_images):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        page_contents[page_num]["images"].append((img_index, image_bytes))
                    except Exception as e:
                        print(f"提取图片失败: {e}")

        doc.close()

        # OCR 识别所有图片（并发处理）
        if enable_ocr:
            # 收集所有需要 OCR 的图片
            ocr_tasks = []
            task_info = []  # 记录每个任务对应的页面和图片索引

            for page_num, content in page_contents.items():
                for img_index, image_bytes in content["images"]:
                    task_info.append((page_num, img_index))
                    ocr_tasks.append(ocr_image_async(image_bytes, vision_provider="aliyun"))

            # 并发执行 OCR
            if ocr_tasks:
                ocr_results = await asyncio.gather(*ocr_tasks, return_exceptions=True)

                # 将 OCR 结果按页面组织
                page_ocr_results = {}
                for i, result in enumerate(ocr_results):
                    page_num, img_index = task_info[i]
                    if page_num not in page_ocr_results:
                        page_ocr_results[page_num] = []

                    if isinstance(result, Exception):
                        print(f"OCR 识别失败（第 {page_num + 1} 页，图片 {img_index + 1}）: {result}")
                    elif result and result.strip():
                        page_ocr_results[page_num].append((img_index, result))

                # 将 OCR 结果存入对应页面
                for page_num, results in page_ocr_results.items():
                    page_contents[page_num]["ocr_results"] = results

        # 组装最终文本，图片内容紧跟在对应页面文本之后
        text_parts = []
        for page_num in sorted(page_contents.keys()):
            content = page_contents[page_num]

            # 添加页面标题和文本
            page_text = content["text"]
            if page_text.strip():
                text_parts.append(f"### 第 {page_num + 1} 页\n{page_text}")

            # 在页面文本之后添加该页面的图片 OCR 结果
            ocr_results = content.get("ocr_results", [])
            for img_index, ocr_text in sorted(ocr_results, key=lambda x: x[0]):
                text_parts.append(f"\n【图片解析开始】（第 {page_num + 1} 页，图片 {img_index + 1}）\n{ocr_text}\n【图片解析结束】\n")

        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 解析失败: {str(e)}")


async def parse_word_to_text(file_bytes: bytes, enable_ocr: bool = True) -> str:
    """
    解析 Word 文档（.docx），提取文本内容、表格和图片（OCR识别）
    图片 OCR 结果会被插入到文档中图片的原始位置，标注为【图片解析】

    Args:
        file_bytes: Word 文件的字节内容
        enable_ocr: 是否启用图片 OCR 识别

    Returns:
        str: 提取的文本内容（包含表格和 OCR 识别的图片内容，位于原始位置）
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import parse_xml
        from io import BytesIO
        import asyncio

        # 辅助函数：正确提取表格行数据，处理合并单元格
        def extract_row_data(row):
            """提取行数据，正确处理合并单元格（避免重复内容）"""
            row_data = []
            processed_cells = set()  # 记录已处理的单元格索引

            for col_idx, cell in enumerate(row.cells):
                # 检查是否已经处理过这个单元格（合并单元格的情况）
                if col_idx in processed_cells:
                    row_data.append("")  # 合并的单元格位置添加空字符串
                    continue

                # 获取单元格属性
                tc_pr = cell._element.tcPr
                if tc_pr is None:
                    # 没有特殊属性，直接添加文本
                    row_data.append(cell.text.strip())
                    continue

                # 检查水平合并 (gridSpan)
                grid_span_element = tc_pr.gridSpan
                if grid_span_element is not None:
                    span_count = int(grid_span_element.get('val', 1))
                    if span_count > 1:
                        # 这是水平合并的起始单元格
                        row_data.append(cell.text.strip())
                        # 标记被合并的列位置
                        for i in range(1, span_count):
                            processed_cells.add(col_idx + i)
                        continue

                # 检查垂直合并 (vMerge) - 使用 find 方法
                v_merge_element = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if v_merge_element is not None:
                    v_merge_val = v_merge_element.get('val', 'continue')
                    if v_merge_val == 'continue':
                        # 这是垂直合并的延续部分
                        row_data.append("")
                        continue
                    else:
                        # 这是垂直合并的起始单元格
                        row_data.append(cell.text.strip())
                        continue

                # 检查水平合并延续 (hMerge) - 使用 find 方法
                h_merge_element = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hMerge')
                if h_merge_element is not None:
                    # 这是水平合并的延续部分
                    row_data.append("")
                    continue

                # 普通单元格
                row_data.append(cell.text.strip())

            return row_data

        # 从字节创建文档
        doc = Document(BytesIO(file_bytes))

        # 1. 首先提取所有图片并建立映射
        image_map = {}  # {rId: image_bytes}
        if enable_ocr:
            for rel_id, rel in doc.part.rels.items():
                if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image":
                    try:
                        image_map[rel_id] = rel.target_part.blob
                        print(f"[Word解析] 发现图片: {rel_id}")
                    except Exception as e:
                        print(f"[Word解析] 提取图片失败 (rId={rel_id}): {e}")

        print(f"[Word解析] 共发现 {len(image_map)} 个图片")

        # 2. 按文档顺序收集内容
        content_items = []  # [(position, type, content)]
        position = 0
        image_tasks = []  # OCR 任务列表
        processed_images = set()  # 已处理的图片 rId

        # 定义命名空间
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        }

        # 遍历文档的 body 元素
        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':  # 段落
                # 提取段落文本
                para_text = ''
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text
                        if para.style and para.style.name:
                            style_name = para.style.name.lower()
                            if 'heading 1' in style_name or 'title' in style_name:
                                para_text = f"# {para.text}"
                            elif 'heading 2' in style_name:
                                para_text = f"## {para.text}"
                            elif 'heading 3' in style_name:
                                para_text = f"### {para.text}"
                        break

                if para_text.strip():
                    content_items.append((position, 'text', para_text))
                    position += 1

                # 查找段落中的图片（使用 XPath 和正确的命名空间）
                if enable_ocr:
                    # 尝试多种方式查找图片引用
                    blips = element.findall('.//a:blip', namespaces)
                    if not blips:
                        # 备用方式：直接搜索包含 embed 属性的元素
                        for child in element.iter():
                            if child.tag.endswith('}blip'):
                                blips.append(child)

                    for blip in blips:
                        # 尝试获取 r:embed 属性
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if not embed_id:
                            embed_id = blip.get(qn('r:embed'))

                        if embed_id and embed_id in image_map and embed_id not in processed_images:
                            img_bytes = image_map[embed_id]
                            processed_images.add(embed_id)
                            content_items.append((position, 'image_placeholder', len(image_tasks)))
                            image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                            position += 1
                            print(f"[Word解析] 段落中发现图片引用: {embed_id}")

            elif tag == 'tbl':  # 表格
                for table in doc.tables:
                    if table._tbl == element:
                        rows_data = []
                        for row in table.rows:
                            row_data = extract_row_data(row)  # 使用新的辅助函数
                            rows_data.append(row_data)

                        if rows_data:
                            header = rows_data[0]
                            table_text = "| " + " | ".join(header) + " |\n"
                            table_text += "| " + " | ".join(["---"] * len(header)) + " |\n"
                            for row in rows_data[1:]:
                                table_text += "| " + " | ".join(row) + " |\n"
                            content_items.append((position, 'table', table_text))
                            position += 1
                        break

        # 3. 处理未被引用但存在的图片（确保所有图片都被处理）
        if enable_ocr:
            unprocessed_images = set(image_map.keys()) - processed_images
            if unprocessed_images:
                print(f"[Word解析] 发现 {len(unprocessed_images)} 个未在段落中引用的图片，追加处理")
                for rel_id in unprocessed_images:
                    img_bytes = image_map[rel_id]
                    content_items.append((position, 'image_placeholder', len(image_tasks)))
                    image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                    position += 1

        # 4. 如果没有提取到任何内容，使用简单备用方式
        if not content_items:
            print("[Word解析] 主要方式未提取到内容，使用备用方式")
            for para in doc.paragraphs:
                if para.text.strip():
                    para_text = para.text
                    if para.style and para.style.name:
                        style_name = para.style.name.lower()
                        if 'heading 1' in style_name or 'title' in style_name:
                            para_text = f"# {para.text}"
                        elif 'heading 2' in style_name:
                            para_text = f"## {para.text}"
                        elif 'heading 3' in style_name:
                            para_text = f"### {para.text}"
                    content_items.append((position, 'text', para_text))
                    position += 1

            for table_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    row_data = extract_row_data(row)  # 使用新的辅助函数处理合并单元格
                    rows_data.append(row_data)
                if rows_data:
                    header = rows_data[0]
                    table_text = f"\n### 表格 {table_idx + 1}\n"
                    table_text += "| " + " | ".join(header) + " |\n"
                    table_text += "| " + " | ".join(["---"] * len(header)) + " |\n"
                    for row in rows_data[1:]:
                        table_text += "| " + " | ".join(row) + " |\n"
                    content_items.append((position, 'table', table_text))
                    position += 1

            # 所有图片追加到末尾
            if enable_ocr and image_map:
                for rel_id, img_bytes in image_map.items():
                    content_items.append((position, 'image_placeholder', len(image_tasks)))
                    image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                    position += 1

        # 5. 并发执行 OCR
        ocr_results = {}
        if image_tasks:
            print(f"[Word解析] 开始 OCR 识别 {len(image_tasks)} 个图片...")
            results = await asyncio.gather(*image_tasks, return_exceptions=True)
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    print(f"[Word解析] 图片 OCR 失败: {result}")
                    ocr_results[i] = None
                elif result and result.strip():
                    ocr_results[i] = result
                    print(f"[Word解析] 图片 {i+1} OCR 成功")
                else:
                    ocr_results[i] = None

        # 6. 组装最终内容
        text_parts = []
        image_counter = 0
        for pos, item_type, content in sorted(content_items, key=lambda x: x[0]):
            if item_type == 'text':
                text_parts.append(content)
            elif item_type == 'table':
                text_parts.append(content)
            elif item_type == 'image_placeholder':
                task_idx = content
                image_counter += 1
                ocr_text = ocr_results.get(task_idx)
                if ocr_text:
                    text_parts.append(f"\n【图片解析开始】（图片 {image_counter}）\n{ocr_text}\n【图片解析结束】\n")

        print(f"[Word解析] 完成，共处理 {image_counter} 个图片")
        return "\n\n".join(text_parts)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise Exception(f"Word 文档解析失败: {str(e)}")


def init_vector_db(file_bytes: bytes, filename: str):
    """
    初始化向量数据库，支持 txt, md, pdf, docx 格式

    Args:
        file_bytes: 文件字节内容
        filename: 文件名

    Returns:
        tuple: (success, persist_dir/error_message, message)
    """
    try:
        from llama_index.core import VectorStoreIndex, Document, Settings
        from llama_index.core.node_parser import SentenceSplitter
        from llama_index.embeddings.huggingface import HuggingFaceEmbedding

        # 根据文件类型解析内容
        file_ext = filename.lower().split('.')[-1]

        if file_ext in ('txt', 'md'):
            text = file_bytes.decode('utf-8')
        elif file_ext == 'pdf':
            text = parse_pdf_to_text(file_bytes)
        elif file_ext in ('docx', 'doc'):
            text = parse_word_to_text(file_bytes)
        else:
            return False, None, f"不支持的文件格式: {file_ext}"

        if not text.strip():
            return False, None, "文档内容为空"

        # 设置全局模型配置
        Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
        Settings.llm = None  # 明确禁用LLM

        # 创建文档
        documents = [Document(text=text, id=filename)]

        # 使用 SentenceSplitter 进行分片（更通用）
        text_splitter = SentenceSplitter(
            chunk_size=512,
            chunk_overlap=50
        )

        # 对文档进行分片
        nodes = text_splitter.get_nodes_from_documents(documents)

        # 创建向量索引
        index = VectorStoreIndex(nodes)

        # 持久化存储
        persist_dir = f"{VECTOR_DB_PATH}/{filename}"
        index.storage_context.persist(persist_dir=persist_dir)

        return True, persist_dir, f"向量数据库初始化完成！文件 {filename} 已成功存储，共 {len(nodes)} 个文本块。"

    except Exception as e:
        return False, None, f"向量数据库初始化失败: {str(e)}"


def query_vector_db_and_call_api(persist_dir, prompt):
    from llama_index.core import StorageContext, load_index_from_storage, Settings
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    # 加载全局配置
    Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)  # 使用指定嵌入模型
    Settings.llm = None  # 显式禁用LLM生成能力

    storage_context = StorageContext.from_defaults(persist_dir=persist_dir)
    index = load_index_from_storage(storage_context)  # 从本地加载预构建的向量索引

    # 配置查询引擎时明确不使用LLM
    query_engine = index.as_query_engine(
        similarity_top_k=3,  # 返回Top3最相似结果
        llm=None  # 关键设置：禁用默认LLM 避免结果被改写 纯向量检索模式：直接返回原始匹配片段，不进行LLM的摘要/重组
    )

    retrieved_docs = query_engine.query(prompt)  # 用用户问题检索
    context = retrieved_docs.response  # 提取原始文本片段
    return context


# def ocr_image(image_path):
#     try:
#         # 打开图片
#         image = Image.open(image_path)
#         # 使用 pytesseract 进行 OCR 识别
#         text = pytesseract.image_to_string(image, lang='chi_sim+eng')  # 支持中文和英文
#         # 将识别结果按行分割成字段
#         fields = text.splitlines()
#         # 过滤掉空行
#         fields = [field.strip() for field in fields if field.strip()]
#         return fields
#     except Exception as e:
#         print(f"处理图片时出错: {e}")
#         return []
# 换行处理函数
# 换行处理函数
async def ocr_image_async(uploaded_files, vision_provider: str = "aliyun", custom_prompt: str = None):
    """
    使用视觉大模型识别图片内容，支持多张图片并发处理（异步版本）
    内置熔断机制：并发限制 / 超时控制 / 失败重试

    Args:
        uploaded_files: 单个图片字节或图片字节列表
        vision_provider: 视觉模型提供商，"aliyun" 或 "deepseek"（默认 aliyun）
        custom_prompt: 自定义识别提示词，为空则使用默认提示词

    Returns:
        str: 识别出的文本，表格会保持结构化格式
    """
    import base64
    from llms import call_vision_api

    # 处理上传的文件对象
    if not uploaded_files:
        return ""
    input_list = [uploaded_files] if isinstance(uploaded_files, bytes) else uploaded_files

    # 使用自定义提示词或默认提示词
    if custom_prompt and custom_prompt.strip():
        prompt = custom_prompt.strip()
    else:
        # 默认表格识别提示词
        prompt = """请识别这张图片中的所有文字内容，并100%还原原始格式和布局。

要求：
1. 如果图片包含表格，必须使用 Markdown 表格格式输出，示例：
   | 列1 | 列2 | 列3 |
   | --- | --- | --- |
   | 数据1 | 数据2 | 数据3 |

2. 表格要求：
   - 完整保留所有行和列，不要遗漏任何单元格
   - 如果有合并单元格，在对应位置重复内容或留空
   - 保持原表格的列数一致
   - 表头和数据行都要完整输出

3. 【重要】如果是非表格内容（如表单、字段列表、按钮、标签等）：
   - 同一行的内容必须保持在同一行，用空格或制表符分隔，绝对不要换行
   - 横向排列的元素必须保持横向排列，例如：字段1  字段2  字段3  字段4
   - 保持原始的相对位置关系，左边的在左边，右边的在右边
   - 不要把横向的内容拆分成多行垂直排列
   - 示例：如果图片中 "姓名" "年龄" "性别" 在同一行，输出应该是：姓名  年龄  性别

4. 只输出识别到的内容，不要添加任何解释、说明或总结

5. 保持原文的语言（中文/英文等），不要翻译"""

    async def process_single_image(idx: int, file_bytes: bytes):
        """处理单张图片（带熔断：信号量+超时+重试）"""
        async with _get_ocr_semaphore():
            for attempt in range(OCR_MAX_RETRIES + 1):
                try:
                    image_base64 = base64.b64encode(file_bytes).decode('utf-8')
                    result = await asyncio.wait_for(
                        call_vision_api(image_base64, prompt, provider=vision_provider),
                        timeout=OCR_TIMEOUT
                    )

                    if "error" in result:
                        if attempt < OCR_MAX_RETRIES:
                            print(f"[OCR] 图片 {idx+1} 第{attempt+1}次失败: {result['error']}，正在重试...")
                            continue
                        return idx, f"图片 {idx + 1}: 识别失败: {result['error']}"

                    choices = result.get("choices", [])
                    if choices:
                        content = choices[0].get("message", {}).get("content", "")
                        if content:
                            return idx, content
                    return idx, f"图片 {idx + 1}: 未识别到文字"
                except asyncio.TimeoutError:
                    if attempt < OCR_MAX_RETRIES:
                        print(f"[OCR] 图片 {idx+1} 第{attempt+1}次超时，正在重试...")
                        continue
                    return idx, f"图片 {idx + 1}: 识别超时"
                except Exception as e:
                    if attempt < OCR_MAX_RETRIES:
                        print(f"[OCR] 图片 {idx+1} 第{attempt+1}次异常: {e}，正在重试...")
                        continue
                    return idx, f"图片 {idx + 1}: 识别处理失败: {str(e)}"

    # 并发处理所有图片（信号量控制并发数）
    tasks = [process_single_image(idx, fb) for idx, fb in enumerate(input_list)]
    results = await asyncio.gather(*tasks)
    # 按原始顺序排序
    results.sort(key=lambda x: x[0])

    # 格式化输出
    all_results = []
    for idx, content in results:
        if len(input_list) > 1:
            all_results.append(f"=== 图片 {idx + 1} ===\n{content}")
        else:
            all_results.append(content)

    final_result = "\n\n".join(all_results)
    print("Vision OCR Result:", final_result)
    return final_result


def _structure_ocr_result(ocr_lines):
    """
    根据 OCR 返回的坐标信息，识别表格结构并格式化输出

    Args:
        ocr_lines: PaddleOCR 返回的识别结果列表，每项为 [[[x1,y1], [x2,y2], [x3,y3], [x4,y4]], (text, confidence)]

    Returns:
        str: 结构化的文本，表格用制表符分隔
    """
    if not ocr_lines:
        return ""

    # 提取每个文本块的信息：文本、中心y坐标、中心x坐标、高度
    text_blocks = []
    for line in ocr_lines:
        box = line[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
        text = line[1][0]  # 文本内容

        # 计算边界框的中心坐标和高度
        y_coords = [point[1] for point in box]
        x_coords = [point[0] for point in box]
        center_y = sum(y_coords) / 4
        center_x = sum(x_coords) / 4
        height = max(y_coords) - min(y_coords)
        min_x = min(x_coords)

        text_blocks.append({
            'text': text,
            'center_y': center_y,
            'center_x': center_x,
            'min_x': min_x,
            'height': height
        })

    if not text_blocks:
        return ""

    # 计算平均行高，用于判断是否在同一行
    avg_height = sum(b['height'] for b in text_blocks) / len(text_blocks)
    row_threshold = avg_height * 0.6  # 行间距阈值

    # 按 y 坐标排序
    text_blocks.sort(key=lambda b: b['center_y'])

    # 分组：将 y 坐标相近的文本块归为同一行
    rows = []
    current_row = [text_blocks[0]]

    for block in text_blocks[1:]:
        # 如果当前块与当前行的 y 坐标差距小于阈值，认为在同一行
        if abs(block['center_y'] - current_row[0]['center_y']) < row_threshold:
            current_row.append(block)
        else:
            rows.append(current_row)
            current_row = [block]

    rows.append(current_row)

    # 判断是否是表格结构（多行且每行有多列）
    multi_col_rows = sum(1 for row in rows if len(row) > 1)
    is_table = multi_col_rows >= 2  # 至少2行有多列才认为是表格

    # 格式化输出
    result_lines = []
    for row in rows:
        # 按 x 坐标排序（从左到右）
        row.sort(key=lambda b: b['min_x'])

        if is_table and len(row) > 1:
            # 表格行：用制表符分隔
            line_text = "\t".join(b['text'] for b in row)
        else:
            # 普通行：用空格连接或直接输出
            line_text = " ".join(b['text'] for b in row)

        result_lines.append(line_text)

    return "\n".join(result_lines)


def _clean_html_text(raw_html: str) -> str:
    try:
        # Remove script/style
        raw_html = re.sub(r"<script[\s\S]*?</script>", " ", raw_html, flags=re.IGNORECASE)
        raw_html = re.sub(r"<style[\s\S]*?</style>", " ", raw_html, flags=re.IGNORECASE)
        # Strip tags
        text = re.sub(r"<[^>]+>", " ", raw_html)
        # Unescape common entities
        text = text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        # Collapse whitespace
        text = re.sub(r"\s+", " ", text)
        return text.strip()
    except Exception:
        return raw_html


def _extract_blue_text_from_html(raw_html: str) -> str:
    """
    从HTML中提取增量需求的文本内容（包括蓝色和红色字体）
    识别多种颜色表示方式：蓝色、红色等用于标识增量需求
    """
    try:
        from bs4 import BeautifulSoup

        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(raw_html, 'html.parser')

        # 扩展增量需求相关的CSS颜色值（包括蓝色和红色）
        increment_colors = [
            # 蓝色系
            'blue', '#0000ff', '#00f', '#0000FF', '#0000ff',
            'rgb(0,0,255)', 'rgb(0, 0, 255)', 'rgb(0,0,255)', 'rgb(0, 0, 255)',
            '#0066cc', '#0066CC', '#3366ff', '#3366FF', '#1e90ff', '#1E90FF',
            '#4169e1', '#4169E1', '#0000CD', '#0000cd', '#191970', '#191970',
            '#000080', '#000080', '#0080FF', '#0080ff', '#0080ff', '#0080FF',
            '#4A90E2', '#4a90e2', '#5B9BD5', '#5b9bd5',
            'rgb(0,102,204)', 'rgb(0,102,255)', 'rgb(51,102,255)', 'rgb(30,144,255)',
            'rgb(65,105,225)', 'rgb(0,0,205)', 'rgb(25,25,112)', 'rgb(0,0,128)',
            'rgb(0,128,255)', 'rgb(74,144,226)', 'rgb(91,155,213)',
            'rgba(0,0,255,', 'rgba(0,102,204,', 'rgba(51,102,255,',
            'dodgerblue', 'royalblue', 'mediumblue', 'darkblue', 'midnightblue',
            'steelblue', 'cornflowerblue', 'lightblue', 'skyblue', 'deepskyblue',

            # 红色系（用于标识增量需求）
            'red', '#ff0000', '#f00', '#FF0000', '#ff0000',
            'rgb(255,0,0)', 'rgb(255, 0, 0)', 'rgb(255,0,0)', 'rgb(255, 0, 0)',
            '#dc143c', '#DC143C', '#b22222', '#B22222', '#8b0000', '#8B0000',
            '#ff4500', '#FF4500', '#ff6347', '#FF6347', '#ff7f50', '#FF7F50',
            'rgb(220,20,60)', 'rgb(178,34,34)', 'rgb(139,0,0)', 'rgb(255,69,0)',
            'rgb(255,99,71)', 'rgb(255,127,80)',
            'rgba(255,0,0,', 'rgba(220,20,60,', 'rgba(178,34,34,',
            'crimson', 'darkred', 'firebrick', 'indianred', 'lightcoral',
            'salmon', 'tomato', 'orangered', 'darkorange'
        ]

        increment_texts = []

        # 查找所有包含增量需求颜色的元素
        for element in soup.find_all(True):
            # 检查style属性
            if element.get('style'):
                style = element.get('style').lower()
                for color in increment_colors:
                    if color.lower() in style:
                        text = element.get_text(strip=True)
                        if text and len(text) > 1:  # 过滤掉太短的文本
                            increment_texts.append(text)
                        break

            # 检查class属性（可能包含增量相关的类名）
            if element.get('class'):
                classes = ' '.join(element.get('class')).lower()
                increment_keywords = ['blue', 'red', 'increment', 'new', 'add', 'modify', 'update',
                                      'change', 'enhance', 'improve', 'feature', 'highlight', 'important']
                if any(keyword in classes for keyword in increment_keywords):
                    text = element.get_text(strip=True)
                    if text and len(text) > 1:
                        increment_texts.append(text)

            # 检查data属性（可能包含颜色信息）
            for attr_name, attr_value in element.attrs.items():
                if isinstance(attr_value, str) and 'color' in attr_name.lower():
                    attr_value_lower = attr_value.lower()
                    for color in increment_colors:
                        if color.lower() in attr_value_lower:
                            text = element.get_text(strip=True)
                            if text and len(text) > 1:
                                increment_texts.append(text)
                            break

        # 使用CSS类提取方法作为补充
        css_increment_texts = _extract_blue_text_from_css_classes(raw_html)
        if css_increment_texts:
            for text in css_increment_texts.split('\n'):
                if text.strip() and text.strip() not in increment_texts:
                    increment_texts.append(text.strip())

        # 使用正则表达式作为补充，查找可能遗漏的增量需求文本
        regex_increment_texts = _extract_blue_text_with_regex(raw_html)
        if regex_increment_texts:
            for text in regex_increment_texts.split('\n'):
                if text.strip() and text.strip() not in increment_texts:
                    increment_texts.append(text.strip())

        # 去重并合并
        unique_texts = []
        seen = set()
        for text in increment_texts:
            if text not in seen:
                seen.add(text)
                unique_texts.append(text)

        return '\n'.join(unique_texts)

    except ImportError:
        # 如果没有BeautifulSoup，使用正则表达式作为备选方案
        return _extract_blue_text_with_regex(raw_html)
    except Exception:
        return ""


def _extract_blue_text_from_css_classes(raw_html: str) -> str:
    """
    从HTML中提取通过CSS类定义的蓝色文本
    只识别真正的蓝色字体，不识别普通文本
    """
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw_html, 'html.parser')
        blue_texts = []

        # 查找所有可能包含蓝色文本的元素
        for element in soup.find_all(True):
            element_id = element.get('id', '')
            element_class = ' '.join(element.get('class', []))

            # 获取元素的文本内容
            text = element.get_text(strip=True)
            if text and len(text) > 1:
                # 只检查元素ID是否匹配已知的蓝色元素模式
                # 基于CSS分析，扩展的蓝色元素ID列表
                blue_element_ids = [
                    'u2838', 'u2839', 'u2840', 'u2841', 'u2871', 'u2872', 'u2873', 'u2874', 'u2875', 'u2876',
                    'u2877', 'u2878', 'u2879', 'u2880', 'u2881', 'u2882', 'u2883', 'u2884', 'u2885', 'u2898',
                    'u2902', 'u2903', 'u2904', 'u2905', 'u2906', 'u2907', 'u2908', 'u2909'
                ]
                if element_id and any(pattern in element_id for pattern in blue_element_ids):
                    blue_texts.append(text)

                # 检查元素类名是否明确包含蓝色相关的关键词
                if any(keyword in element_class.lower() for keyword in
                       ['blue', 'highlight', 'new', 'add', 'modify', 'increment']):
                    blue_texts.append(text)

        # 去重
        unique_texts = []
        seen = set()
        for text in blue_texts:
            if text not in seen:
                seen.add(text)
                unique_texts.append(text)

        return '\n'.join(unique_texts)

    except ImportError:
        # 如果没有BeautifulSoup，使用正则表达式作为备选
        return _extract_blue_text_from_css_classes_regex(raw_html)
    except Exception:
        return ""


def _extract_blue_text_from_css_classes_regex(raw_html: str) -> str:
    """
    使用正则表达式提取CSS类定义的蓝色文本的备选方案
    只识别真正的蓝色字体，不识别普通文本
    """
    try:
        blue_texts = []

        # 扩展的蓝色元素ID列表（基于CSS分析结果）
        blue_element_ids = [
            'u2838', 'u2839', 'u2840', 'u2841', 'u2871', 'u2872', 'u2873', 'u2874', 'u2875', 'u2876',
            'u2877', 'u2878', 'u2879', 'u2880', 'u2881', 'u2882', 'u2883', 'u2884', 'u2885', 'u2898',
            'u2902', 'u2903', 'u2904', 'u2905', 'u2906', 'u2907', 'u2908', 'u2909'
        ]

        # 生成ID匹配模式
        id_patterns = []
        for element_id in blue_element_ids:
            # 复杂模式：匹配嵌套HTML结构
            id_patterns.append(rf'<[^>]*id="{element_id}"[^>]*>.*?<[^>]*>([^<]+)</[^>]*>.*?</[^>]*>')
            # 简化模式：匹配ID元素内的所有文本
            id_patterns.append(rf'<[^>]*id="{element_id}"[^>]*>([^<]*)</[^>]*>')

        for pattern in id_patterns:
            matches = re.findall(pattern, raw_html, re.IGNORECASE | re.DOTALL)
            blue_texts.extend(matches)

        # 匹配明确包含蓝色相关类名的元素
        class_patterns = [
            r'<[^>]*class="[^"]*blue[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*class="[^"]*highlight[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*class="[^"]*new[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*class="[^"]*increment[^"]*"[^>]*>([^<]+)</[^>]*>',
        ]

        for pattern in class_patterns:
            matches = re.findall(pattern, raw_html, re.IGNORECASE)
            blue_texts.extend(matches)

        # 去重
        unique_texts = []
        seen = set()
        for text in blue_texts:
            cleaned_text = text.strip()
            if cleaned_text and cleaned_text not in seen:
                seen.add(cleaned_text)
                unique_texts.append(cleaned_text)

        return '\n'.join(unique_texts)

    except Exception:
        return ""


def _extract_blue_text_with_regex(raw_html: str) -> str:
    """
    使用正则表达式提取增量需求文本的备选方案（包括蓝色和红色）
    """
    try:
        # 扩展的增量需求匹配模式（包括蓝色和红色）
        increment_patterns = [
            # 标准蓝色
            r'<[^>]*style="[^"]*color\s*:\s*(?:blue|#0000ff|#00f|#0000FF)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # RGB格式蓝色
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*0\s*,\s*0\s*,\s*255\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*0\s*,\s*102\s*,\s*204\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*51\s*,\s*102\s*,\s*255\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*30\s*,\s*144\s*,\s*255\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*65\s*,\s*105\s*,\s*225\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 十六进制蓝色变体
            r'<[^>]*style="[^"]*color\s*:\s*#0066cc[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#0066FF[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#3366ff[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#1e90ff[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#4169e1[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#0000CD[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#000080[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#0080FF[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#4A90E2[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#5B9BD5[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 带透明度的蓝色
            r'<[^>]*style="[^"]*color\s*:\s*rgba\s*\(\s*0\s*,\s*0\s*,\s*255\s*,[^)]*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgba\s*\(\s*0\s*,\s*102\s*,\s*204\s*,[^)]*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 命名蓝色
            r'<[^>]*style="[^"]*color\s*:\s*(?:dodgerblue|royalblue|mediumblue|darkblue|midnightblue|steelblue|cornflowerblue|lightblue|skyblue|deepskyblue)[^"]*"[^>]*>([^<]+)</[^>]*>',

            # 标准红色
            r'<[^>]*style="[^"]*color\s*:\s*(?:red|#ff0000|#f00|#FF0000)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # RGB格式红色
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*255\s*,\s*0\s*,\s*0\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*220\s*,\s*20\s*,\s*60\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*178\s*,\s*34\s*,\s*34\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgb\s*\(\s*139\s*,\s*0\s*,\s*0\s*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 十六进制红色变体
            r'<[^>]*style="[^"]*color\s*:\s*#dc143c[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#DC143C[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#b22222[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#B22222[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#8b0000[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#8B0000[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#ff4500[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#FF4500[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#ff6347[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#FF6347[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#ff7f50[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*#FF7F50[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 带透明度的红色
            r'<[^>]*style="[^"]*color\s*:\s*rgba\s*\(\s*255\s*,\s*0\s*,\s*0\s*,[^)]*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgba\s*\(\s*220\s*,\s*20\s*,\s*60\s*,[^)]*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            r'<[^>]*style="[^"]*color\s*:\s*rgba\s*\(\s*178\s*,\s*34\s*,\s*34\s*,[^)]*\)[^"]*"[^>]*>([^<]+)</[^>]*>',
            # 命名红色
            r'<[^>]*style="[^"]*color\s*:\s*(?:crimson|darkred|firebrick|indianred|lightcoral|salmon|tomato|orangered|darkorange)[^"]*"[^>]*>([^<]+)</[^>]*>',

            # 单引号样式
            r"<[^>]*style='[^']*color\s*:\s*(?:blue|red|#0000ff|#ff0000|#0000FF|#FF0000|#0066cc|#0066FF|#3366ff|#1e90ff|#4169e1|#dc143c|#DC143C)[^']*'[^>]*>([^<]+)</[^>]*>",
            # 无引号样式
            r'<[^>]*style=[^>]*color\s*:\s*(?:blue|red|#0000ff|#ff0000|#0000FF|#FF0000|#0066cc|#0066FF|#3366ff|#1e90ff|#4169e1|#dc143c|#DC143C)[^>]*>([^<]+)</[^>]*>',
        ]

        increment_texts = []
        for pattern in increment_patterns:
            matches = re.findall(pattern, raw_html, re.IGNORECASE)
            increment_texts.extend(matches)

        # 去重
        unique_texts = []
        seen = set()
        for text in increment_texts:
            cleaned_text = text.strip()
            if cleaned_text and cleaned_text not in seen:
                seen.add(cleaned_text)
                unique_texts.append(cleaned_text)

        return '\n'.join(unique_texts)

    except Exception:
        return ""


def _extract_axure_js_strings(js_text: str) -> List[str]:
    extracted: List[str] = []
    try:
        # Common Axure key fields
        patterns = [
            r"\bname\s*:\s*\"([^\"]+)\"",
            r"\blabel\s*:\s*\"([^\"]+)\"",
            r"\btype\s*:\s*\"([^\"]+)\"",
            r"\bnotes\s*:\s*\{[\s\S]*?\btext\s*:\s*\"([\s\S]*?)\"[\s\S]*?\}",
            r"\btext\s*:\s*\{[\s\S]*?\bexpr\b[\s\S]*?\}|\btext\s*:\s*\"([\s\S]*?)\"",
            r"\btip\s*:\s*\"([\s\S]*?)\"",
        ]
        for pat in patterns:
            for m in re.finditer(pat, js_text):
                for g in m.groups():
                    if g:
                        cleaned = re.sub(r"\\n|\\r", " ", g)
                        cleaned = re.sub(r"\s+", " ", cleaned).strip()
                        if cleaned:
                            extracted.append(cleaned)
    except Exception:
        pass
    return extracted


def _collect_blue_selectors_from_css(css_text: str) -> Dict[str, set]:
    """
    从CSS文本中收集设置为蓝色(color)的选择器，返回包含id与class名称的集合。
    仅匹配明确设置为蓝色系的规则，避免误报。
    """
    blue_color_patterns = [
        r"blue\b", r"#0000ff\b", r"#00f\b", r"rgb\s*\(\s*0\s*,\s*0\s*,\s*255\s*\)",
        r"#0066cc\b", r"#3366ff\b", r"#1e90ff\b", r"#4169e1\b"
    ]
    color_regex = r"(?:" + "|".join(blue_color_patterns) + r")"

    blue_ids: set = set()
    blue_classes: set = set()

    try:
        # 形如: #u2838 { color:#0000FF; }
        for m in re.finditer(r"#([\w-]+)\s*\{[^}]*?color\s*:\s*" + color_regex + r"[^}]*\}", css_text, re.IGNORECASE):
            blue_ids.add(m.group(1))

        # 形如: #u2838 .text { color:#0000FF; }
        for m in re.finditer(r"#([\\w-]+)\s+[^\{]*\{[^}]*?color\s*:\s*" + color_regex + r"[^}]*\}", css_text,
                             re.IGNORECASE):
            blue_ids.add(m.group(1))

        # 形如: .blue-text { color: rgb(0,0,255); }
        for m in re.finditer(r"\.([\w-]+)\s*\{[^}]*?color\s*:\s*" + color_regex + r"[^}]*\}", css_text, re.IGNORECASE):
            blue_classes.add(m.group(1))
    except Exception:
        pass

    return {"ids": blue_ids, "classes": blue_classes}


def _extract_text_by_selectors(html: str, blue_ids: set, blue_classes: set) -> List[str]:
    """
    使用已知蓝色id与class集合，从HTML中提取对应元素的文本。
    """
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        results: List[str] = []

        for bid in blue_ids:
            el = soup.find(id=bid)
            if el:
                txt = el.get_text(strip=True)
                if txt:
                    results.append(txt)

        if blue_classes:
            # 任意包含这些类名的元素
            class_selector = ",".join([f".{c}" for c in blue_classes])
            for el in soup.select(class_selector):
                txt = el.get_text(strip=True)
                if txt:
                    results.append(txt)
        return results
    except Exception:
        return []


def parse_axure_zip_to_text(file_bytes: bytes) -> Dict[str, str]:
    """
    Parse an Axure HTML export zipped package and return both full content and incremental content.
    - Extracts text from HTML files
    - Scrapes useful strings from Axure JS data files
    - Separates full content and blue text (incremental requirements)

    Returns:
        Dict with keys: 'full_content', 'incremental_content'
    """
    full_texts: List[str] = []
    incremental_texts: List[str] = []

    try:
        # 在遍历前，先收集所有CSS中声明为蓝色的选择器，提升召回
        css_blue_ids: set = set()
        css_blue_classes: set = set()
        with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
            # 第一遍扫描：聚合CSS蓝色选择器
            for info in zf.infolist():
                name_lower = info.filename.lower()
                if name_lower.endswith(".css") and not name_lower.startswith("__macosx/"):
                    with zf.open(info, "r") as fp:
                        raw = fp.read()
                        try:
                            css_text = raw.decode("utf-8", errors="ignore")
                        except Exception:
                            css_text = raw.decode(errors="ignore")
                        selectors = _collect_blue_selectors_from_css(css_text)
                        css_blue_ids.update(selectors.get("ids", set()))
                        css_blue_classes.update(selectors.get("classes", set()))

            # 第二遍：解析HTML/JS并结合CSS选择器抽取蓝色文本
            for info in zf.infolist():
                name_lower = info.filename.lower()
                if name_lower.endswith((".html", ".htm")) and not name_lower.startswith("__macosx/"):
                    with zf.open(info, "r") as fp:
                        raw = fp.read()
                        try:
                            html = raw.decode("utf-8", errors="ignore")
                        except Exception:
                            html = raw.decode(errors="ignore")

                        # 提取全部内容
                        full_text = _clean_html_text(html)
                        if full_text:
                            full_texts.append(full_text)

                        # 提取蓝色文本（增量需求）
                        blue_text = _extract_blue_text_from_html(html)
                        # 结合CSS选择器再做一次提取，弥补遗漏
                        if css_blue_ids or css_blue_classes:
                            more = _extract_text_by_selectors(html, css_blue_ids, css_blue_classes)
                            if more:
                                blue_text = (blue_text + "\n" + "\n".join(more)).strip() if blue_text else "\n".join(
                                    more)
                        if blue_text:
                            incremental_texts.append(blue_text)

                elif name_lower.endswith(".js") and (
                        "data" in name_lower or "pages" in name_lower or "document" in name_lower):
                    with zf.open(info, "r") as fp:
                        raw = fp.read()
                        try:
                            js_text = raw.decode("utf-8", errors="ignore")
                        except Exception:
                            js_text = raw.decode(errors="ignore")
                        js_strings = _extract_axure_js_strings(js_text)
                        full_texts.extend(js_strings)

    except zipfile.BadZipFile:
        raise Exception("Axure包不是有效的ZIP文件")
    except Exception as e:
        raise Exception(f"解析Axure ZIP失败: {str(e)}")

    # 处理全部内容
    full_merged = "\n".join([t for t in full_texts if t])
    full_parts = []
    seen_full = set()
    for line in full_merged.split("\n"):
        key = line.strip()
        if key and key not in seen_full:
            seen_full.add(key)
            full_parts.append(key)

    # 处理增量内容
    incremental_merged = "\n".join([t for t in incremental_texts if t])
    incremental_parts = []
    seen_incremental = set()
    for line in incremental_merged.split("\n"):
        key = line.strip()
        if key and key not in seen_incremental:
            seen_incremental.add(key)
            incremental_parts.append(key)

    return {
        'full_content': "\n".join(full_parts),
        'incremental_content': "\n".join(incremental_parts)
    }


def parse_axure_html_to_text(file_bytes: bytes) -> Dict[str, str]:
    """
    Parse a single Axure HTML file and return both full content and incremental content.

    Returns:
        Dict with keys: 'full_content', 'incremental_content'
    """
    try:
        html = file_bytes.decode("utf-8", errors="ignore")

        # 提取全部内容
        full_content = _clean_html_text(html)

        # 提取蓝色文本（增量需求）
        incremental_content = _extract_blue_text_from_html(html)

        return {
            'full_content': full_content,
            'incremental_content': incremental_content
        }
    except Exception as e:
        raise Exception(f"解析Axure HTML失败: {str(e)}")


def format_axure_text_to_markdown(axure_text: str) -> str:
    """
    Convert cleaned Axure extracted text to lightweight Markdown structure.
    Heuristics: promote likely page/component names to headings, others to bullet points.
    """
    if not axure_text:
        return ""
    lines = [l.strip() for l in axure_text.split("\n") if l and l.strip()]
    md_lines: List[str] = ["# 原型需求提取\n"]
    last_was_heading = False
    for line in lines:
        # heading if short and title-like
        if (len(line) <= 40 and
                re.search(r"[：:]|(页面|功能|模块|流程|用例|说明|规则|字段)$", line) or
                (line.istitle() and not re.search(r"\s", line) and len(line) <= 20)):
            if not last_was_heading:
                md_lines.append("")
            md_lines.append(f"## {line}")
            last_was_heading = True
        else:
            # turn into bullet point; split potential key-value
            bullet = line
            kv = re.split(r"[：:]", line, maxsplit=1)
            if len(kv) == 2 and len(kv[0]) <= 20:
                bullet = f"**{kv[0].strip()}**: {kv[1].strip()}"
            md_lines.append(f"- {bullet}")
            last_was_heading = False
    return "\n".join(md_lines).strip()


def format_incremental_text_to_markdown(incremental_text: str) -> str:
    """
    Convert incremental (blue text) content to Markdown structure.
    """
    if not incremental_text:
        return ""

    lines = [l.strip() for l in incremental_text.split("\n") if l and l.strip()]
    if not lines:
        return ""

    md_lines: List[str] = ["# 增量需求提取\n"]
    md_lines.append("> 以下内容为原型中的蓝色字体部分，通常表示新增或修改的需求\n\n")

    for line in lines:
        # 检查是否是标题类型的内容
        if (len(line) <= 50 and
                (re.search(r"[：:]|(新增|修改|更新|优化|改进|功能|模块|页面|流程)$", line) or
                 line.istitle() and len(line) <= 30)):
            md_lines.append(f"## {line}")
        else:
            # 作为列表项处理
            md_lines.append(f"- {line}")

    return "\n".join(md_lines).strip()


def refine_requirements_markdown(requirement_text: str, provider: str | None = None) -> str:
    """
    Use LLM to refine and structure requirement text into strict Markdown sections.
    Returns Markdown string.
    同步版本 - 用于非async上下文
    """
    import asyncio

    prompt = f"""
你是一名文档整理专家。请将以下从原型提取的文字整理为清晰的Markdown格式，要求：

1. **保持原型的原始结构和层级关系**，不要改变内容的组织方式

2. **表格处理（重要）**：
   - 如果内容包含表格数据（如用 | 分隔的内容，或者有明显的行列结构），必须转换为标准Markdown表格
   - 表格格式示例：
     ```
     | 列1 | 列2 | 列3 | 列4 |
     |-----|-----|-----|-----|
     | 数据1 | 数据2 | 数据3 | 数据4 |
     ```
   - 表头行和数据行之间必须有 `|-----|-----|` 分隔线
   - 确保每行的列数一致
   - 如果原文有类似"模块|功能|类型|说明"这样的结构，识别为表格并格式化

3. 使用合适的Markdown标记：
   - 标题用 # ## ###
   - 列表用 - 或 1. 2. 3.
   - 重要内容可以用 **加粗**
   - 代码或技术术语用 `反引号`

4. 清理格式问题：
   - 去除多余的空行和空格
   - 修复破碎的句子
   - 合并重复的内容

5. **不要添加原文没有的内容**，不要臆造字段或功能

6. **不要改变原型的结构**，如果原型是按页面/模块组织的，保持这种组织方式

7. 如果有"参看原型"、"详情参看"等引用，保留这些引用关系

待整理内容：
{requirement_text}
"""

    try:
        resp = asyncio.run(call_llm_api(prompt, provider=provider))
        content = resp.get("choices", [{}])[0].get("message", {}).get("content", "")
        return content.strip() if content else requirement_text
    except Exception:
        return requirement_text


async def refine_requirements_markdown_async(requirement_text: str, provider: str | None = None) -> str:
    """
    Use LLM to refine and structure requirement text into strict Markdown sections.
    Returns Markdown string.
    异步版本 - 用于async上下文（如FastAPI）
    """
    print(f"[DEBUG] refine_requirements_markdown_async 被调用, provider={provider}, 文本长度={len(requirement_text)}")

    prompt = f"""
你是一名文档整理专家。请将以下从原型提取的文字整理为清晰的Markdown格式，要求：

1. **保持原型的原始结构和层级关系**，不要改变内容的组织方式，去掉菜单功能描述，只保留页面主要功能描述和介绍

2. **表格处理（重要）**：
   - 如果内容包含表格数据，必须转换为标准 Markdown 表格格式
   - 标准格式示例：
     ```
     | 列1 | 列2 | 列3 |
     |-----|-----|-----|
     | 数据1 | 数据2 | 数据3 |
     ```
   - 确保每行的列数一致，表头和分隔符完整
   - 识别"字段|类型|说明"等结构并格式化为表格

3. **流程图处理（重要）**：
   - 如果内容描述了流程、步骤、状态转换，请用 Mermaid 流程图格式输出
   - 格式示例：
     ```mermaid
     flowchart TD
         A[开始] --> B{{判断条件}}
         B -->|是| C[执行操作1]
         B -->|否| D[执行操作2]
         C --> E[结束]
         D --> E
     ```
   - 状态流转用 `stateDiagram-v2`
   - 时序图用 `sequenceDiagram`

4. 使用合适的Markdown标记：
   - 标题用 # ## ###
   - 列表用 - 或 1. 2. 3.
   - 重要内容可以用 **加粗**
   - 代码或技术术语用 `反引号`

5. 清理格式问题：
   - 去除多余的空行和空格
   - 修复破碎的句子
   - 合并重复的内容

6. **不要添加原文没有的内容**，不要臆造字段或功能

7. **不要改变原型的结构**，保持按页面/模块组织的方式

8. 保留"参看原型"、"详情参看"等引用关系

9.你只需要返回待整理的内容即可 不要返回其他内容

待整理内容：
{requirement_text}
"""

    try:
        print(f"[DEBUG] 开始调用LLM API...")
        resp = await call_llm_api(prompt, provider=provider)
        print(f"[DEBUG] LLM API调用完成, 响应: {str(resp)[:200]}...")
        content = resp.get("choices", [{}])[0].get("message", {}).get("content", "")
        if content:
            print(f"[DEBUG] AI结构化成功, 结果长度={len(content)}")
            return content.strip()
        else:
            print(f"[DEBUG] AI返回内容为空，使用原始文本")
            return requirement_text
    except Exception as e:
        print(f"[DEBUG] AI结构化失败: {e}")
        import traceback
        traceback.print_exc()
        return requirement_text


# ==================== Chroma + 通义千问Embedding 向量化功能 ====================

def get_dashscope_embedding(texts: List[str]) -> List[List[float]]:
    """
    使用通义千问 text-embedding-v3 模型获取文本向量

    Args:
        texts: 文本列表

    Returns:
        List[List[float]]: 向量列表
    """
    import dashscope
    from dashscope import TextEmbedding

    # 从环境变量获取API Key
    api_key = os.getenv("ALIYUN_API_KEY") or os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise Exception("缺少 ALIYUN_API_KEY 或 DASHSCOPE_API_KEY 环境变量")

    dashscope.api_key = api_key

    embeddings = []

    # 逐条处理，避免批量大小限制问题
    for i, text in enumerate(texts):
        print(f"正在向量化第 {i+1}/{len(texts)} 个文本块...")
        response = TextEmbedding.call(
            model=DASHSCOPE_EMBEDDING_MODEL,
            input=text  # 单条文本，不是列表
        )

        if response.status_code != 200:
            raise Exception(f"Embedding API调用失败: {response.message}")

        # 单条返回的结构
        embedding = response.output['embeddings'][0]['embedding']
        embeddings.append(embedding)

    return embeddings


def split_text_into_chunks(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    将文本切割成块

    Args:
        text: 原始文本
        chunk_size: 每块的最大字符数
        chunk_overlap: 块之间的重叠字符数

    Returns:
        List[str]: 文本块列表
    """
    if not text or not text.strip():
        return []

    # 按段落分割
    paragraphs = re.split(r'\n\s*\n', text)

    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 如果当前段落本身就超过chunk_size，需要进一步切割
        if len(para) > chunk_size:
            # 先保存当前chunk
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""

            # 按句子切割长段落
            sentences = re.split(r'([。！？；\n])', para)
            temp_chunk = ""

            for i in range(0, len(sentences), 2):
                sentence = sentences[i]
                if i + 1 < len(sentences):
                    sentence += sentences[i + 1]  # 加上标点

                if len(temp_chunk) + len(sentence) <= chunk_size:
                    temp_chunk += sentence
                else:
                    if temp_chunk:
                        chunks.append(temp_chunk.strip())
                    # 保留重叠部分
                    if chunk_overlap > 0 and len(temp_chunk) > chunk_overlap:
                        temp_chunk = temp_chunk[-chunk_overlap:] + sentence
                    else:
                        temp_chunk = sentence

            if temp_chunk:
                current_chunk = temp_chunk
        else:
            # 正常段落处理
            if len(current_chunk) + len(para) + 1 <= chunk_size:
                current_chunk += "\n" + para if current_chunk else para
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                # 保留重叠部分
                if chunk_overlap > 0 and len(current_chunk) > chunk_overlap:
                    current_chunk = current_chunk[-chunk_overlap:] + "\n" + para
                else:
                    current_chunk = para

    # 添加最后一个chunk
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


async def init_chroma_vector_db(file_bytes: bytes, filename: str, base_path: str = None) -> tuple:
    """
    使用 Chroma + 通义千问Embedding 初始化向量数据库
    支持 txt, md, pdf, docx 格式
    自动使用阿里云 OCR 识别文档中的图片
    MD 文件中的图片引用会被解析并还原到原位置

    Args:
        file_bytes: 文件字节内容
        filename: 文件名
        base_path: 文件所在的基础路径（用于解析 MD 文件中的相对路径图片）

    Returns:
        tuple: (success, collection_name/error_message, message)
    """
    try:
        import chromadb
        from chromadb.config import Settings

        # 确保目录存在
        os.makedirs(CHROMA_DB_PATH, exist_ok=True)

        # 根据文件类型解析内容
        file_ext = filename.lower().split('.')[-1]

        if file_ext == 'txt':
            # TXT 文件直接解码，无图片处理
            text = file_bytes.decode('utf-8')
        elif file_ext == 'md':
            # MD 文件先解码，然后解析其中的图片引用
            text = file_bytes.decode('utf-8')
            # 解析 MD 中的图片并还原到原位置
            text = await parse_markdown_with_images(text, base_path=base_path, vision_provider="aliyun")
            print(f"[MD向量化] 文件 {filename} 图片解析完成")
        elif file_ext == 'pdf':
            text = await parse_pdf_to_text(file_bytes, enable_ocr=True)
        elif file_ext in ('docx', 'doc'):
            text = await parse_word_to_text(file_bytes, enable_ocr=True)
        else:
            return False, None, f"不支持的文件格式: {file_ext}"

        if not text.strip():
            return False, None, "文档内容为空"

        # 切割文本
        chunks = split_text_into_chunks(text, chunk_size=500, chunk_overlap=50)

        if not chunks:
            return False, None, "文档切割后无有效内容"

        print(f"文档 {filename} 切割成 {len(chunks)} 个文本块")

        # 获取文本向量
        embeddings = get_dashscope_embedding(chunks)

        # 创建Chroma客户端 (持久化存储)
        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)

        # 生成唯一的collection名称 (基于文件名和时间戳)
        timestamp = get_current_datetime()
        # collection名称只能包含字母数字和下划线，长度3-63
        safe_filename = re.sub(r'[^a-zA-Z0-9]', '_', filename.rsplit('.', 1)[0])[:30]
        collection_name = f"doc_{safe_filename}_{timestamp}"

        # 创建或获取collection
        collection = client.get_or_create_collection(
            name=collection_name,
            metadata={
                "filename": filename,
                "created_at": timestamp,
                "original_document": text  # 保存原始完整文档内容
            }
        )

        # 添加文档到collection
        ids = [f"chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"chunk_index": i, "filename": filename} for i in range(len(chunks))]

        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas
        )

        return True, collection_name, f"向量数据库初始化完成！文件 {filename} 已成功存储到 Chroma，共 {len(chunks)} 个文本块。"

    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, None, f"Chroma向量数据库初始化失败: {str(e)}"


def query_chroma_vector_db(collection_name: str, query_text: str, top_k: int = 5, recall_all: bool = False) -> Dict[str, Any]:
    """
    从 Chroma 向量数据库中检索相关内容

    Args:
        collection_name: collection名称
        query_text: 查询文本（全量召回时可以为空）
        top_k: 返回的最相似结果数量
        recall_all: 是否全量召回（返回集合中所有文档）

    Returns:
        Dict: 包含检索结果的字典
    """
    try:
        import chromadb

        # 连接到持久化的Chroma数据库
        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)

        # 获取collection
        try:
            collection = client.get_collection(name=collection_name)
        except Exception:
            return {"success": False, "error": f"未找到collection: {collection_name}"}

        # 全量召回：返回所有文档块并按顺序合并
        if recall_all:
            # 获取 collection 中所有文档
            all_data = collection.get(include=["documents", "metadatas"])

            documents = all_data.get("documents", [])
            metadatas = all_data.get("metadatas", [])

            if not documents:
                return {"success": False, "error": "集合中没有文档内容"}

            # 按 chunk_index 排序，确保顺序正确
            doc_with_meta = list(zip(documents, metadatas))
            doc_with_meta.sort(key=lambda x: x[1].get("chunk_index", 0) if x[1] else 0)

            # 合并所有文档块
            sorted_docs = [doc for doc, _ in doc_with_meta]
            full_content = "\n\n".join(sorted_docs)

            print(f"[全量召回] 共 {len(documents)} 个文档块，合并后长度: {len(full_content)}")

            return {
                "success": True,
                "context": full_content,
                "chunks": [{"rank": i+1, "content": doc, "distance": None, "metadata": meta}
                          for i, (doc, meta) in enumerate(doc_with_meta)],
                "total_chunks": len(documents),
                "is_original": True
            }
        else:
            # 正常检索：只返回 top_k 个结果
            # 获取查询文本的向量
            query_embedding = get_dashscope_embedding([query_text])[0]
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )

        # 整理结果（正常检索模式）
        documents = results.get('documents', [[]])[0]
        distances = results.get('distances', [[]])[0]
        metadatas = results.get('metadatas', [[]])[0]

        retrieved_chunks = []
        for i, (doc, dist, meta) in enumerate(zip(documents, distances, metadatas)):
            retrieved_chunks.append({
                "rank": i + 1,
                "content": doc,
                "distance": dist,
                "metadata": meta
            })

        # 合并检索到的内容作为上下文
        context = "\n\n---\n\n".join([chunk["content"] for chunk in retrieved_chunks])

        return {
            "success": True,
            "context": context,
            "chunks": retrieved_chunks,
            "total_chunks": len(retrieved_chunks)
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": f"Chroma向量检索失败: {str(e)}"}


def list_chroma_collections() -> List[Dict[str, Any]]:
    """
    列出所有可用的Chroma collections

    Returns:
        List[Dict]: collection列表
    """
    try:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        collections = client.list_collections()

        result = []
        for col in collections:
            result.append({
                "name": col.name,
                "metadata": col.metadata,
                "count": col.count()
            })

        return result

    except Exception as e:
        print(f"列出collections失败: {str(e)}")
        return []


def check_document_exists(filename: str) -> tuple:
    """
    检查文档是否已经向量化存储

    Args:
        filename: 文件名

    Returns:
        tuple: (exists: bool, collection_name: str or None)
    """
    try:
        collections = list_chroma_collections()
        for col in collections:
            # 检查 metadata 中的 filename 是否匹配
            if col.get("metadata", {}).get("filename") == filename:
                return True, col.get("name")
        return False, None
    except Exception:
        return False, None


def delete_chroma_collection(collection_name: str) -> bool:
    """
    删除指定的Chroma collection

    Args:
        collection_name: collection名称

    Returns:
        bool: 是否删除成功
    """
    try:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        client.delete_collection(name=collection_name)
        return True

    except Exception as e:
        print(f"删除collection失败: {str(e)}")
        return False


# ==================== 流程图检测与转换 ====================

def detect_and_extract_flowchart(soup, iframe_content: str) -> Dict[str, Any]:
    """
    检测并提取Axure页面中的流程图元素

    Args:
        soup: BeautifulSoup对象
        iframe_content: 原始HTML内容

    Returns:
        Dict包含:
        - has_flowchart: 是否检测到流程图
        - nodes: 节点列表 [{text, x, y, width, height, shape_type}]
        - description: 流程图的文字描述
    """
    result = {
        'has_flowchart': False,
        'nodes': [],
        'connections': [],
        'description': ''
    }

    # 检测流程图的特征
    # 1. Axure 使用 img 标签引用外部 SVG 作为连接线/箭头
    # 特征: src 包含 _seg0.svg, _seg1.svg 等模式，或者在 images 目录下的 SVG
    has_connector_images = False
    connector_count = 0

    for img in soup.find_all('img'):
        src = img.get('src', '')
        img_id = img.get('id', '')
        # 检测连接线图片: _seg 模式 或 arrow/connector 相关
        if ('_seg' in src.lower() or '_seg' in img_id.lower() or
            'arrow' in src.lower() or 'connector' in src.lower() or
            (src.endswith('.svg') and '/images/' in src)):
            connector_count += 1
            if connector_count >= 2:  # 至少2个连接线
                has_connector_images = True

    # 2. 也检查内嵌 SVG 元素（备用）
    svg_elements = soup.find_all('svg')
    has_svg_lines = False
    for svg in svg_elements:
        if svg.find_all(['path', 'line', 'polyline']):
            has_svg_lines = True
            break

    # 3. 检测带有位置信息的形状元素
    positioned_elements = []
    flowchart_keywords = ['开始', '结束', '判断', '是', '否', 'yes', 'no', 'start', 'end',
                          '流程', '步骤', '条件', '分支', '循环', '处理', '输入', '输出',
                          '提交', '审核', '审批', '通过', '拒绝', '完成', '发起', '申请']

    # 查找所有绝对定位的div元素
    for div in soup.find_all('div', style=True):
        style = div.get('style', '')
        # 检查是否是绝对定位且有位置信息
        if 'position' in style.lower() and ('left' in style.lower() or 'top' in style.lower()):
            text = div.get_text(strip=True)
            if text and len(text) < 100:  # 流程图节点文字通常较短
                # 解析位置信息
                left_match = re.search(r'left:\s*(-?\d+(?:\.\d+)?)\s*px', style, re.IGNORECASE)
                top_match = re.search(r'top:\s*(-?\d+(?:\.\d+)?)\s*px', style, re.IGNORECASE)
                width_match = re.search(r'width:\s*(-?\d+(?:\.\d+)?)\s*px', style, re.IGNORECASE)
                height_match = re.search(r'height:\s*(-?\d+(?:\.\d+)?)\s*px', style, re.IGNORECASE)

                if left_match and top_match:
                    node = {
                        'text': text,
                        'x': float(left_match.group(1)),
                        'y': float(top_match.group(1)),
                        'width': float(width_match.group(1)) if width_match else 100,
                        'height': float(height_match.group(1)) if height_match else 50,
                        'shape_type': 'process'  # 默认为处理框
                    }

                    # 根据关键词判断形状类型
                    text_lower = text.lower()
                    if any(k in text_lower for k in ['开始', 'start', '起始', '发起']):
                        node['shape_type'] = 'start'
                    elif any(k in text_lower for k in ['结束', 'end', '完成', '终止']):
                        node['shape_type'] = 'end'
                    elif any(k in text_lower for k in ['判断', '条件', '是否', '?', '？', '审核', '审批']):
                        node['shape_type'] = 'decision'
                    elif text in ['是', '否', 'yes', 'no', 'Y', 'N', '通过', '拒绝', '同意', '不同意']:
                        node['shape_type'] = 'label'  # 连接线上的标签

                    positioned_elements.append(node)

    # 判断是否是流程图
    # 条件：有多个定位元素 + (有连接线图片 或 有SVG连接线 或 有流程图关键词)
    has_flowchart_keywords = any(
        any(kw in node['text'].lower() for kw in flowchart_keywords)
        for node in positioned_elements
    )

    if len(positioned_elements) >= 3 and (has_connector_images or has_svg_lines or has_flowchart_keywords):
        result['has_flowchart'] = True
        result['nodes'] = positioned_elements
        result['connector_count'] = connector_count

        # 按位置排序（先上后下，先左后右）
        sorted_nodes = sorted(positioned_elements, key=lambda n: (n['y'], n['x']))

        # 生成描述
        descriptions = []
        for i, node in enumerate(sorted_nodes):
            if node['shape_type'] != 'label':
                descriptions.append(f"{i+1}. [{node['shape_type']}] {node['text']}")

        result['description'] = '\n'.join(descriptions)
        print(f"[流程图检测] 检测到 {len(positioned_elements)} 个节点, {connector_count} 个连接线图片")

    return result


async def convert_flowchart_to_mermaid_async(flowchart_data: Dict[str, Any], provider: str = "deepseek") -> str:
    """
    调用AI将流程图描述转换为Mermaid代码（异步版本）

    Args:
        flowchart_data: detect_and_extract_flowchart返回的数据
        provider: LLM提供商

    Returns:
        Mermaid格式的流程图代码
    """
    if not flowchart_data.get('has_flowchart') or not flowchart_data.get('nodes'):
        return ""

    # 构建提示词
    nodes_desc = []
    for node in flowchart_data['nodes']:
        if node['shape_type'] != 'label':
            nodes_desc.append(f"- 类型:{node['shape_type']}, 文字:「{node['text']}」, 位置:(x={node['x']}, y={node['y']})")

    labels = [n for n in flowchart_data['nodes'] if n['shape_type'] == 'label']
    labels_desc = ', '.join([f"「{l['text']}」" for l in labels]) if labels else "无"

    prompt = f"""请根据以下从Axure原型中提取的流程图元素，生成Mermaid格式的流程图代码。

## 提取的节点信息（按位置排序，y值越小越靠上）:
{chr(10).join(nodes_desc)}

## 连接线上的标签文字:
{labels_desc}

## 要求:
1. 使用 flowchart TD（从上到下）或 flowchart LR（从左到右）格式
2. 根据节点的位置关系（y值）推断连接顺序
3. 开始节点使用圆角矩形 ([文字])
4. 结束节点使用圆角矩形 ([文字])
5. 判断/条件节点使用菱形 {{文字}}
6. 普通处理节点使用矩形 [文字]
7. 如果有"是/否"等标签，添加到连接线上
8. 只输出Mermaid代码，不要其他解释

## 示例输出格式:
```mermaid
flowchart TD
    A([开始]) --> B[处理步骤1]
    B --> C{{判断条件}}
    C -->|是| D[处理步骤2]
    C -->|否| E[处理步骤3]
    D --> F([结束])
    E --> F
```

请生成Mermaid代码:"""

    try:
        response = await call_llm_api(prompt, provider=provider)
        content = response.get("choices", [{}])[0].get("message", {}).get("content", "")

        if not content:
            return ""

        # 提取代码块中的内容
        if '```mermaid' in content:
            match = re.search(r'```mermaid\s*([\s\S]*?)\s*```', content)
            if match:
                return match.group(1).strip()
        elif '```' in content:
            match = re.search(r'```\s*([\s\S]*?)\s*```', content)
            if match:
                return match.group(1).strip()

        # 如果没有代码块，检查是否直接返回了flowchart代码
        if content.strip().startswith('flowchart'):
            return content.strip()

        return ""
    except Exception as e:
        print(f"[流程图转换] AI转换失败: {e}")
        return ""


# ==================== Axure 在线链接获取功能 ====================

def fetch_axure_from_url(url: str, username: str = None, password: str = None,
                          wait_time: int = 5) -> Dict[str, str]:
    """
    使用无头浏览器获取Axure在线原型内容（同步版本）

    优化实现：更好地解析 Axure 页面结构，提取表格和结构化内容

    Args:
        url: Axure在线原型链接
        username: 域账号用户名（可选）
        password: 域账号密码（可选）
        wait_time: 页面加载等待时间（秒）

    Returns:
        Dict with keys: 'full_content', 'incremental_content'
    """
    try:
        from playwright.sync_api import sync_playwright
        from playwright_stealth import Stealth
        from bs4 import BeautifulSoup
        from urllib.parse import urlparse, unquote
        import time

        with sync_playwright() as p:
            # 启动无头浏览器
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(
                viewport={'width': 1920, 'height': 1080},
                user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            )
            page = context.new_page()

            # 使用stealth避免被检测
            stealth = Stealth()
            stealth.apply_stealth_sync(page)

            print(f"正在访问Axure链接: {url}")

            # 访问页面
            page.goto(url, wait_until='networkidle', timeout=60000)
            time.sleep(wait_time)

            # 处理登录
            if username and password:
                page_title = page.title()
                if '认证' in page_title or 'login' in page_title.lower():
                    print("检测到登录页面，正在登录...")
                    try:
                        login_methods = [
                            ('input[placeholder="请输入您的用户名"]', 'input[placeholder="请输入您的密码"]'),
                            ('input[name="username"]', 'input[name="password"]'),
                            ('input[type="text"]', 'input[type="password"]'),
                        ]

                        for user_sel, pwd_sel in login_methods:
                            try:
                                if page.locator(user_sel).count() > 0:
                                    page.locator(user_sel).fill(username)
                                    page.locator(pwd_sel).fill(password)
                                    break
                            except Exception:
                                continue

                        login_btns = ['button:has-text("登录")', 'input[type="submit"]', 'button[type="submit"]']
                        for btn_sel in login_btns:
                            try:
                                if page.locator(btn_sel).count() > 0:
                                    page.locator(btn_sel).click()
                                    break
                            except Exception:
                                continue

                        page.wait_for_load_state('networkidle')
                        time.sleep(3)
                    except Exception as login_error:
                        print(f"登录处理警告: {login_error}")

            # 从 URL 解析页面名称
            parsed_url = urlparse(url)
            page_name = ""
            if 'p=' in url:
                import urllib.parse
                query_params = urllib.parse.parse_qs(parsed_url.query)
                if 'p' in query_params:
                    page_name = unquote(query_params['p'][0])

            time.sleep(2)

            # ============= 获取Axure主内容区域iframe =============
            print(f"[Axure解析] 开始获取主内容区域...")

            iframe_content = ""

            # Axure典型的iframe名称/id列表
            main_frame_selectors = [
                'iframe#mainFrame',
                'iframe[name="mainFrame"]',
                'iframe#mainPanel',
                'iframe[name="mainPanel"]',
                '#mainFrame',
                '[name="mainFrame"]',
                'iframe.mainFrame',
            ]

            # 方法1: 尝试通过选择器定位主内容iframe
            for selector in main_frame_selectors:
                try:
                    frame_element = page.locator(selector)
                    if frame_element.count() > 0:
                        # 等待iframe加载
                        frame_element.wait_for(state='attached', timeout=5000)
                        # 获取frame对象
                        frame = page.frame(name='mainFrame') or page.frame(url=lambda u: 'mainFrame' in str(u) or '.html' in str(u))
                        if frame:
                            iframe_content = frame.content()
                            print(f"[Axure解析] 通过选择器 {selector} 获取到mainFrame内容: {len(iframe_content)} 字符")
                            break
                except Exception as e:
                    continue

            # 方法2: 遍历所有frames，查找包含实际内容的frame（排除导航frame）
            if not iframe_content or len(iframe_content) < 500:
                print(f"[Axure解析] 选择器方式未找到，尝试遍历所有frames...")
                frames = page.frames
                print(f"[Axure解析] 共发现 {len(frames)} 个frames")

                best_frame_content = ""
                for i, frame in enumerate(frames):
                    try:
                        frame_url = frame.url
                        frame_name = frame.name
                        content = frame.content()
                        content_len = len(content)

                        print(f"[Axure解析] Frame {i}: name={frame_name}, url={frame_url[:50] if frame_url else ''}..., 内容长度={content_len}")

                        # 跳过主frame和明显的导航frame
                        if frame == page.main_frame:
                            continue

                        # 跳过sitemap/导航相关的frame
                        frame_url_lower = frame_url.lower() if frame_url else ''
                        frame_name_lower = frame_name.lower() if frame_name else ''
                        if any(keyword in frame_url_lower or keyword in frame_name_lower
                               for keyword in ['sitemap', 'toc', 'nav', 'menu', 'tree', 'left']):
                            print(f"[Axure解析] 跳过导航frame: {frame_name}")
                            continue

                        # 选择内容最长的非导航frame
                        if content_len > len(best_frame_content):
                            best_frame_content = content
                            print(f"[Axure解析] 更新最佳frame内容: {content_len} 字符")

                    except Exception as e:
                        print(f"[Axure解析] 获取frame {i} 内容失败: {e}")
                        continue

                if best_frame_content:
                    iframe_content = best_frame_content

            # 方法3: 如果还是没有内容，获取整个页面
            if not iframe_content:
                print(f"[Axure解析] 未找到iframe，使用整页内容")
                iframe_content = page.content()

            print(f"[Axure解析] 最终获取内容长度: {len(iframe_content)}")

            # 保存target_frame用于后续获取computed color
            target_frame = page
            # 查找主内容frame
            for frame in page.frames:
                frame_url = frame.url or ''
                frame_name = frame.name or ''
                # 跳过导航frame
                if any(kw in frame_url.lower() or kw in frame_name.lower()
                       for kw in ['sitemap', 'toc', 'nav', 'menu', 'tree', 'left']):
                    continue
                if frame != page.main_frame and len(frame_url) > 10:
                    target_frame = frame
                    break

            # 使用BeautifulSoup解析
            soup = BeautifulSoup(iframe_content, 'html.parser')

            full_texts = []
            incremental_texts = []
            tables_data = []

            # 移除脚本和样式
            for tag in soup(['script', 'style', 'noscript', 'link', 'meta']):
                tag.decompose()

            # 过滤关键词 - Axure UI 界面元素
            axure_ui_keywords = [
                # 基础UI
                'preview', 'inspect', 'share', 'adaptive', 'comments', 'hotspots',
                'collapse all', 'scale to', 'default scale', 'user scale',
                'show note markers', 'copyright', 'axure', 'prototype',
                'close', 'variables', 'zoom', 'pages', 'masters', 'console',
                # Inspect 面板
                'colors', 'assets', 'size and position', 'download all',
                'copied to clipboard', 'typography', 'typeface', 'fill color',
                'border', 'shadows', 'no notes for this page', 'add comment',
                'mark all read', 'rotation', 'radius', 'padding', 'opacity',
                'width:', 'height:', 'align:', 'position:', 'size:',
                # 右侧面板
                'add a comment', 'give feedback', 'ask a question', 'request a change',
                # 其他
                'sitemap', 'outline', 'notes', 'interactions', 'documentation',
                'publish', 'generate', 'export', 'import', 'settings'
            ]

            # 完全匹配过滤（单独的短词）
            axure_exact_keywords = {'content', 'html', 'css', 'other'}

            code_keywords = [
                'function(', 'var ', 'const ', 'let ', 'return ', 'console.',
                'jquery', 'document.', 'window.', '$(', 'css(', 'axure.',
                '{', '}', '===', '!=='
            ]

            def extract_text_smart(element):
                text = element.get_text(strip=True)
                if not text or len(text) < 2:
                    return None
                text_lower = text.lower().strip()
                # 完全匹配过滤
                if text_lower in axure_exact_keywords:
                    return None
                # 检查是否包含 Axure UI 关键词
                if any(keyword in text_lower for keyword in axure_ui_keywords):
                    return None
                if any(keyword in text for keyword in code_keywords):
                    return None
                # 过滤纯数字
                if text.isdigit():
                    return None
                # 过滤类似 "X:" "Y:" "Width:" 等短标签
                if len(text) <= 10 and text.endswith(':'):
                    return None
                # 过滤组合标签如 "X:Y:" "Width:Height:"
                if re.match(r'^[A-Za-z:]+$', text) and ':' in text:
                    return None
                # 检查是否包含有效字符
                if not any(c.isalpha() or c > '\u4e00' for c in text):
                    return None
                if text.isascii() and len(text) < 3:
                    return None
                # 过滤重复内容如 "×Colors" "×Assets"
                if text.startswith('×'):
                    return None
                return text

            # 提取表格数据
            for table in soup.find_all('table'):
                rows = []
                for tr in table.find_all('tr'):
                    cells = []
                    for td in tr.find_all(['td', 'th']):
                        cell_text = td.get_text(strip=True)
                        cells.append(cell_text)
                    if cells and any(c for c in cells):
                        rows.append(cells)
                if rows:
                    tables_data.append(rows)

            # 提取文本元素
            seen_texts = set()

            def normalize_style_sync(style_str):
                """标准化样式字符串，去除所有空格并转小写"""
                return style_str.lower().replace(' ', '').replace('\t', '').replace('\n', '')

            def check_color_in_style_sync(style_str):
                """检查样式字符串中是否包含蓝色/红色"""
                normalized = normalize_style_sync(style_str)
                blue_patterns_normalized = [
                    'color:blue', 'color:#0000ff', 'color:#00f',
                    'color:rgb(0,0,255)', 'color:#0066cc', 'color:#0066ff',
                    'color:#3366ff', 'color:#1e90ff', 'color:#4169e1',
                    'color:#0000cd', 'color:#000080', 'color:#0080ff',
                    'color:#4a90e2', 'color:#5b9bd5', 'color:#2196f3',
                    'color:#1976d2', 'color:dodgerblue', 'color:royalblue',
                    'color:steelblue', 'color:cornflowerblue',
                    'color:red', 'color:#ff0000', 'color:#f00',
                    'color:rgb(255,0,0)', 'color:#dc143c', 'color:#ff4500',
                    'color:#ff6347', 'color:#b22222', 'color:#8b0000',
                    'color:#cc0000', 'color:crimson', 'color:darkred',
                    'color:firebrick',
                ]
                for pattern in blue_patterns_normalized:
                    if pattern in normalized:
                        return True
                # 使用正则匹配 rgb 格式
                if re.search(r'color[:\s]*rgb\s*\(\s*0\s*,\s*0\s*,\s*255\s*\)', style_str, re.IGNORECASE):
                    return True
                if re.search(r'color[:\s]*rgb\s*\(\s*255\s*,\s*0\s*,\s*0\s*\)', style_str, re.IGNORECASE):
                    return True
                return False

            def check_element_color_sync(element):
                """检查元素及其祖先是否有蓝色/红色样式"""
                # 检查元素自身
                style = element.get('style', '')
                if style and check_color_in_style_sync(style):
                    return True
                # 检查父元素（最多5层）
                parent = element.parent
                for _ in range(5):
                    if parent is None or parent.name is None:
                        break
                    parent_style = parent.get('style', '')
                    if parent_style and check_color_in_style_sync(parent_style):
                        return True
                    parent = parent.parent
                return False

            def check_cell_has_blue_color(td_element):
                """检查单元格内是否有蓝色/红色内容（包括子元素）"""
                # 检查单元格自身
                if check_element_color_sync(td_element):
                    return True
                # 检查单元格内的所有子元素
                for child in td_element.find_all(recursive=True):
                    if check_element_color_sync(child):
                        return True
                return False

            # ============= 使用JavaScript获取computed color =============
            # 这是关键修复：Axure渲染后的蓝色是通过CSS设置的，不是inline style
            # 所以需要用JavaScript获取computedStyle.color
            # 注意：这段代码必须在browser.close()之前执行
            blue_texts_from_js = set()
            try:
                print(f"[Axure解析] 使用JavaScript获取computed color...")
                color_info = target_frame.evaluate("""
                () => {
                    const results = [];
                    const elements = document.querySelectorAll('div, span, p, h1, h2, h3, h4, h5, h6, li, a, td, th, label');
                    for (const el of elements) {
                        const text = el.innerText ? el.innerText.trim() : '';
                        if (text && text.length > 1 && text.length < 500) {
                            const style = window.getComputedStyle(el);
                            const color = style.color;
                            results.push({
                                text: text.substring(0, 200),
                                color: color
                            });
                        }
                    }
                    return results;
                }
                """)

                if color_info:
                    for info in color_info:
                        color = info.get('color', '')
                        text = info.get('text', '').strip()
                        if not text or not color:
                            continue
                        # 检查是否是蓝色系或红色系
                        is_blue = False
                        if 'rgb' in color:
                            try:
                                parts = color.replace('rgb(', '').replace('rgba(', '').replace(')', '').split(',')
                                r, g, b = int(parts[0].strip()), int(parts[1].strip()), int(parts[2].strip())
                                # 蓝色: B高(>150), R低(<150)
                                if b > 150 and r < 150:
                                    is_blue = True
                                # 红色: R高(>200), G和B低(<100)
                                elif r > 200 and g < 100 and b < 100:
                                    is_blue = True
                            except:
                                pass
                        if is_blue:
                            blue_texts_from_js.add(text)
                            # 也添加文本的子串（因为元素可能嵌套，导致文本重复）
                            for line in text.split('\n'):
                                line = line.strip()
                                if line and len(line) > 1:
                                    blue_texts_from_js.add(line)

                print(f"[Axure解析] JavaScript识别到 {len(blue_texts_from_js)} 条蓝色/红色文本")
            except Exception as js_err:
                print(f"[Axure解析] JavaScript获取computed color失败: {js_err}")

            def is_text_blue_by_js(text):
                """检查文本是否被JavaScript识别为蓝色"""
                if not text:
                    return False
                text = text.strip()
                # 精确匹配
                if text in blue_texts_from_js:
                    return True
                # 检查文本是否是某个蓝色文本的子串（文本被包含在蓝色文本中）
                # 注意：不能反过来检查，否则会把包含蓝色子串的长文本都标记为蓝色
                for blue_text in blue_texts_from_js:
                    if text in blue_text:
                        return True
                return False
            # ============= JavaScript computed color 检测结束 =============

            # 提取表格数据 - 同时检测蓝色行
            incremental_table_rows = []  # 存储包含蓝色内容的表格行
            table_headers = []  # 存储表头

            for table in soup.find_all('table'):
                rows = []
                current_header = []
                for row_idx, tr in enumerate(table.find_all('tr')):
                    cells = []
                    row_has_blue = False  # 标记该行是否有蓝色内容

                    for td in tr.find_all(['td', 'th']):
                        cell_text = td.get_text(strip=True)
                        cell_text = ' '.join(cell_text.split())  # 清理空白
                        cells.append(cell_text)

                        # 检查该单元格是否有蓝色/红色内容
                        # 方法1: 检查inline style（原有逻辑）
                        if check_cell_has_blue_color(td):
                            row_has_blue = True
                        # 方法2: 检查JavaScript computed color（新增逻辑）
                        elif is_text_blue_by_js(cell_text):
                            row_has_blue = True

                    if cells and any(c for c in cells):
                        rows.append(cells)

                        # 第一行通常是表头
                        if row_idx == 0:
                            current_header = cells
                        elif row_has_blue and current_header:
                            # 如果该行有蓝色内容，记录整行（带表头信息）
                            incremental_table_rows.append({
                                'header': current_header,
                                'row': cells
                            })

                if rows:
                    tables_data.append(rows)
                    if current_header:
                        table_headers.append(current_header)

            # 提取非表格的文本元素
            seen_texts = set()

            for element in soup.find_all(['div', 'span', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'a']):
                # 跳过表格内的元素（已单独处理）
                if element.find_parent('table'):
                    continue

                text = extract_text_smart(element)
                if not text:
                    continue

                # 检查是否是增量内容（蓝色/红色）
                # 方法1: 检查inline style（原有逻辑）
                is_increment = check_element_color_sync(element)
                # 方法2: 检查JavaScript computed color（新增逻辑，修复CSS设置颜色的情况）
                if not is_increment:
                    is_increment = is_text_blue_by_js(text)

                # 增量内容：即使文本重复也要记录
                if is_increment and text not in incremental_texts:
                    incremental_texts.append(text)

                # 全量内容：去重
                if text not in seen_texts:
                    seen_texts.add(text)
                    full_texts.append(text)

            browser.close()

            # 构建 Markdown 输出 - 全量内容
            md_parts = []

            if page_name:
                md_parts.append(f"# {page_name}\n")

            if tables_data:
                for idx, table_rows in enumerate(tables_data):
                    if len(table_rows) > 1:
                        max_cols = max(len(row) for row in table_rows)
                        md_parts.append(f"\n## 表格 {idx + 1}\n")
                        header = table_rows[0] if table_rows else []
                        header_padded = header + [''] * (max_cols - len(header))
                        md_parts.append("| " + " | ".join(header_padded) + " |")
                        md_parts.append("| " + " | ".join(['---'] * max_cols) + " |")
                        for row in table_rows[1:]:
                            row_padded = row + [''] * (max_cols - len(row))
                            row_cleaned = [cell.replace('\n', ' ').replace('\r', '').replace('|', '｜') for cell in row_padded]
                            md_parts.append("| " + " | ".join(row_cleaned) + " |")
                        md_parts.append("")

            if full_texts:
                md_parts.append("\n## 页面内容\n")
                for text in full_texts[:200]:
                    if len(text) <= 50 and re.search(r'[：:](|页面|功能|模块|流程|说明|规则|字段|管理|配置)$', text):
                        md_parts.append(f"\n### {text}\n")
                    else:
                        md_parts.append(f"- {text}")

            full_content = "\n".join(md_parts)
            full_content = re.sub(r'\n{3,}', '\n\n', full_content)

            # 构建增量内容 - 包含蓝色表格行和蓝色文本
            inc_parts = []

            # 1. 先输出包含蓝色内容的表格行（以完整表格格式）
            if incremental_table_rows:
                inc_parts.append("# 增量需求（蓝色/红色标记内容）\n")
                inc_parts.append("## 增量表格字段\n")

                # 按表头分组显示
                header_groups = {}
                for item in incremental_table_rows:
                    header_key = tuple(item['header'])
                    if header_key not in header_groups:
                        header_groups[header_key] = []
                    header_groups[header_key].append(item['row'])

                for header, rows in header_groups.items():
                    header_list = list(header)
                    max_cols = max(len(header_list), max(len(row) for row in rows))

                    # 输出表头
                    header_padded = header_list + [''] * (max_cols - len(header_list))
                    inc_parts.append("| " + " | ".join(header_padded) + " |")
                    inc_parts.append("| " + " | ".join(['---'] * max_cols) + " |")

                    # 输出蓝色行
                    for row in rows:
                        row_padded = row + [''] * (max_cols - len(row))
                        row_cleaned = [cell.replace('\n', ' ').replace('\r', '').replace('|', '｜') for cell in row_padded]
                        inc_parts.append("| " + " | ".join(row_cleaned) + " |")

                    inc_parts.append("")  # 表格间空行

            # 2. 再输出非表格的蓝色文本
            if incremental_texts:
                if not inc_parts:
                    inc_parts.append("# 增量需求（蓝色/红色标记内容）\n")
                if incremental_table_rows:
                    inc_parts.append("## 其他增量内容\n")
                for text in incremental_texts:
                    inc_parts.append(f"- {text}")

            incremental_content = "\n".join(inc_parts) if inc_parts else ""

            print(f"提取完成: 全量内容 {len(full_content)} 字符, 增量内容 {len(incremental_content)} 字符")
            print(f"提取了 {len(tables_data)} 个表格, {len(full_texts)} 条文本, {len(incremental_table_rows)} 条增量表格行")

            # 检测流程图
            flowchart_data = detect_and_extract_flowchart(soup, iframe_content)
            if flowchart_data.get('has_flowchart'):
                print(f"[流程图检测] 检测到流程图，共 {len(flowchart_data.get('nodes', []))} 个节点")

            return {
                'full_content': full_content.strip(),
                'incremental_content': incremental_content.strip(),
                'flowchart_data': flowchart_data  # 包含流程图检测结果
            }

    except ImportError as e:
        raise Exception(f"playwright库未安装，请先运行: pip install playwright && playwright install chromium")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise Exception(f"获取Axure在线内容失败: {str(e)}")


def _extract_links_from_blue_content(html_content: str, base_url: str) -> List[Dict[str, str]]:
    """
    从HTML内容中提取蓝色/红色文本中的链接（包括文字描述的页面引用）
    [保留此函数作为备用]
    """
    # 直接使用文本提取函数
    return _extract_page_references_from_text(html_content, base_url)


def _extract_page_references_from_text(text_content: str, base_url: str) -> List[Dict[str, str]]:
    """
    从文本内容中提取页面引用（如"参看xxx页面"）

    Args:
        text_content: 文本内容（可以是增量内容或HTML）
        base_url: 基础URL，用于构建完整链接

    Returns:
        List[Dict]: 包含链接信息的列表 [{'url': '...', 'text': '...', 'page_name': '...'}]
    """
    from urllib.parse import urlparse, unquote
    import re

    links = []
    seen_page_names = set()

    try:
        # 解析基础URL
        parsed_base = urlparse(base_url)
        # 构建基础URL（不含查询参数）
        base_url_without_query = f"{parsed_base.scheme}://{parsed_base.netloc}{parsed_base.path}"

        # 页面引用的关键词模式 - 更宽泛的匹配
        reference_patterns = [
            # "参看xxx" 系列
            r'参看[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?',
            r'详见[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?',
            r'参见[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?',
            r'见[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?[页原]',
            # "跳转到xxx" 系列
            r'跳转[到至]?[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?',
            r'链接[到至]?[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?',
            # "查看xxx" 系列
            r'查看[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?[页原]',
            r'参考[「「\[【《]?([^」」\]】》\s，,。.；;：:]+)[」」\]】》]?[页原]',
            # 直接引用格式
            r'[（\(]参[看见][「「\[【《]?([^」」\]】》\s，,。.；;：:\)）]+)[」」\]】》]?[）\)]?',
            r'[（\(]详见[「「\[【《]?([^」」\]】》\s，,。.；;：:\)）]+)[」」\]】》]?[）\)]?',
        ]

        for pattern in reference_patterns:
            matches = re.findall(pattern, text_content)
            for match in matches:
                page_name = match.strip()
                if not page_name or len(page_name) < 2:
                    continue

                # 清理页面名称
                page_name = page_name.replace('页面', '').replace('原型', '').replace('页', '').strip()
                # 去除可能的标点符号
                page_name = re.sub(r'^[「「\[【《]+|[」」\]】》]+$', '', page_name).strip()

                if not page_name or len(page_name) < 2:
                    continue

                # 跳过已处理的页面
                if page_name in seen_page_names:
                    continue
                seen_page_names.add(page_name)

                # 构建URL
                full_url = f"{base_url_without_query}?p={page_name}"

                links.append({
                    'url': full_url,
                    'text': f"引用: {page_name}",
                    'page_name': page_name
                })
                print(f"[页面引用提取] 找到: '{page_name}' -> {full_url}")

        print(f"[页面引用提取] 共提取到 {len(links)} 个页面引用")
        return links

    except Exception as e:
        print(f"[页面引用提取] 提取失败: {e}")
        import traceback
        traceback.print_exc()
        return []


async def fetch_axure_from_url_async_recursive(
    url: str,
    username: str = None,
    password: str = None,
    wait_time: int = 2,
    max_depth: int = 3,
    enable_recursive: bool = True
) -> Dict[str, Any]:
    """
    递归版本：获取Axure在线原型内容，并递归解析蓝色内容中的链接

    Args:
        url: Axure原型链接
        username: 登录用户名
        password: 登录密码
        wait_time: 页面等待时间
        max_depth: 最大递归深度（默认3层）
        enable_recursive: 是否启用递归解析

    Returns:
        Dict with keys:
        - 'full_content': 主页面全量内容
        - 'incremental_content': 按层级组织的增量内容
        - 'pages': 每个页面的单独内容列表（新增）
    """
    from urllib.parse import urlparse, unquote, parse_qs
    import asyncio

    visited_urls = set()
    all_results = []  # 存储所有页面的结果
    pages_list = []  # 新增：存储每个页面的完整信息

    def get_heading_prefix(depth: int) -> str:
        """根据深度返回Markdown标题前缀"""
        if depth == 1:
            return "#"
        elif depth == 2:
            return "##"
        elif depth == 3:
            return "###"
        else:
            return "####"  # 超过3层用4级标题

    async def fetch_page_recursive(page_url: str, depth: int):
        """递归获取页面内容"""
        # 检查是否已访问或超过深度限制
        if page_url in visited_urls:
            print(f"[递归解析] 跳过已访问的页面: {page_url[:50]}...")
            return
        if depth > max_depth:
            print(f"[递归解析] 已达最大深度 {max_depth}，停止递归")
            return

        visited_urls.add(page_url)

        # 解析页面名称
        parsed = urlparse(page_url)
        query_params = parse_qs(parsed.query)
        page_name = ""
        if 'p' in query_params:
            page_name = unquote(query_params['p'][0])

        print(f"[递归解析] 深度={depth}/{max_depth}, 页面={page_name or page_url[:50]}...")

        try:
            # 获取当前页面内容
            result = await fetch_axure_from_url_async(
                url=page_url,
                username=username,
                password=password,
                wait_time=wait_time
            )

            full_content = result.get('full_content', '')
            incremental_content = result.get('incremental_content', '')
            raw_html = result.get('raw_html', '')

            # 存储当前页面的完整信息（新增）
            pages_list.append({
                'page_name': page_name or f"页面{len(pages_list) + 1}",
                'page_url': page_url,
                'full_content': full_content,
                'incremental_content': incremental_content,
                'depth': depth,
                'has_incremental': bool(incremental_content.strip())
            })

            # 存储当前页面结果（用于构建合并的增量内容）
            if incremental_content:
                # 清理原有标题
                clean_content = incremental_content.replace('# 增量需求（蓝色/红色标记内容）\n', '').strip()
                # 移除列表项的前导符号，保持原始文本
                clean_lines = []
                for line in clean_content.split('\n'):
                    if line.startswith('- '):
                        clean_lines.append(line[2:])
                    else:
                        clean_lines.append(line)
                all_results.append({
                    'depth': depth,
                    'page_name': page_name or f"页面{len(all_results) + 1}",
                    'content': '\n'.join(clean_lines)
                })

            # 如果启用递归且未达最大深度，从增量内容中提取页面引用
            if enable_recursive and depth < max_depth and incremental_content:
                # 从已提取的增量内容文字中查找页面引用
                links = _extract_page_references_from_text(incremental_content, page_url)
                print(f"[递归解析] 从增量内容中提取到 {len(links)} 个页面引用")

                # 递归获取每个链接的内容
                for link_info in links:
                    link_url = link_info['url']
                    if link_url not in visited_urls:
                        try:
                            await fetch_page_recursive(link_url, depth + 1)
                        except Exception as link_error:
                            print(f"[递归解析] 获取链接失败，跳过: {link_url[:50]}... 错误: {link_error}")
                            continue

        except Exception as e:
            print(f"[递归解析] 获取页面失败 ({page_url[:50]}...): {e}")
            return

    # 开始递归获取
    print(f"[递归解析] 开始递归解析，最大深度={max_depth}, 启用递归={enable_recursive}")
    await fetch_page_recursive(url, depth=1)

    # 如果不启用递归或没有结果，至少添加主页面
    if not pages_list:
        result = await fetch_axure_from_url_async(
            url=url,
            username=username,
            password=password,
            wait_time=wait_time
        )
        from urllib.parse import urlparse, parse_qs, unquote
        parsed = urlparse(url)
        query_params = parse_qs(parsed.query)
        page_name = unquote(query_params.get('p', ['主页面'])[0])

        pages_list.append({
            'page_name': page_name,
            'page_url': url,
            'full_content': result.get('full_content', ''),
            'incremental_content': result.get('incremental_content', ''),
            'depth': 1,
            'has_incremental': bool(result.get('incremental_content', '').strip())
        })

    # 构建带层级结构的增量内容
    final_incremental_parts = []

    for item in all_results:
        depth = item['depth']
        page_name = item['page_name']
        content = item['content']

        heading_prefix = get_heading_prefix(depth)

        # 添加带层级的标题
        final_incremental_parts.append(f"\n{heading_prefix} {page_name}\n")

        # 添加内容（作为列表项）
        for line in content.split('\n'):
            line = line.strip()
            if line:
                final_incremental_parts.append(f"- {line}")

    final_incremental_content = "\n".join(final_incremental_parts).strip()

    # 获取主页面的全量内容
    main_result = await fetch_axure_from_url_async(
        url=url,
        username=username,
        password=password,
        wait_time=wait_time
    )

    print(f"[递归解析] 完成! 共解析 {len(pages_list)} 个页面")

    return {
        'full_content': main_result.get('full_content', ''),
        'incremental_content': final_incremental_content,
        'pages': pages_list  # 新增：每个页面的单独内容
    }


async def fetch_axure_from_url_async(url: str, username: str = None, password: str = None,
                                     wait_time: int = 2, provider: str = "deepseek") -> Dict[str, str]:
    """
    异步版本：使用无头浏览器获取Axure在线原型内容

    通过线程池运行同步版本，避免 Windows 上 asyncio 与 Playwright 的兼容性问题
    如果检测到流程图，会自动转换为 Mermaid 格式并插入到内容中
    """
    import asyncio
    from functools import partial

    # 使用 asyncio.to_thread 在线程池中运行同步版本
    # 这样可以避免 Windows 上的事件循环冲突问题
    func = partial(fetch_axure_from_url, url, username, password, wait_time)
    result = await asyncio.to_thread(func)

    # 处理流程图：如果检测到流程图，转换为 Mermaid 并插入内容
    flowchart_data = result.get('flowchart_data', {})
    if flowchart_data.get('has_flowchart'):
        print(f"[流程图处理] 开始将流程图转换为 Mermaid...")
        try:
            mermaid_code = await convert_flowchart_to_mermaid_async(flowchart_data, provider=provider)
            if mermaid_code:
                print(f"[流程图处理] Mermaid 转换成功，代码长度: {len(mermaid_code)}")
                # 不在这里添加到 full_content，由调用方决定何时添加
                # 这样可以避免被格式化函数破坏
                result['mermaid_code'] = mermaid_code  # 单独返回 Mermaid 代码供前端使用
            else:
                print(f"[流程图处理] Mermaid 转换返回空结果")
        except Exception as e:
            print(f"[流程图处理] Mermaid 转换失败: {e}")

    return result


# ==================== 增强版知识库管理功能 ====================

# 知识库存储路径
KNOWLEDGE_BASE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "knowledge_bases")

# 确保知识库目录存在
os.makedirs(KNOWLEDGE_BASE_PATH, exist_ok=True)


# ============ 元数据Schema定义 ============

METADATA_SCHEMA = {
    "doc_type": "page_full | page_incremental | module_fused | business_flow",
    "page_key": "页面唯一标识",
    "page_name": "页面中文名称",
    "module": "所属模块",
    "business_domain": "业务域",
    "is_incremental": "Boolean: 是否为增量内容",
    "has_blue_text": "Boolean: 是否包含蓝色字体",
    "related_pages": "关联的页面keys列表",
    "prerequisite_pages": "前置页面",
    "subsequent_pages": "后续页面",
    "content_length": "内容长度",
    "created_at": "创建时间",
}


# ============ 数据模型 ============

class PageData:
    """页面数据模型"""
    def __init__(self, page_key: str, page_name: str, page_url: str, full_content: str = "",
                 incremental_content: str = "", module: str = "", raw_html: str = ""):
        self.page_key = page_key           # 页面唯一标识
        self.page_name = page_name          # 页面名称
        self.page_url = page_url            # 页面URL
        self.module = module                # 所属模块（AI分析后填充）
        self.full_content = full_content    # 全量内容（去除蓝色字体）
        self.incremental_content = incremental_content  # 增量内容（仅蓝色字体）
        self.has_incremental = bool(incremental_content.strip())  # 是否有增量
        self.related_pages = []             # 关联页面（AI分析后填充）
        self.prerequisite_pages = []        # 前置页面
        self.subsequent_pages = []          # 后续页面
        self.raw_html = raw_html            # HTML原始内容

    def to_dict(self) -> dict:
        """转换为字典"""
        return {
            "page_key": self.page_key,
            "page_name": self.page_name,
            "page_url": self.page_url,
            "module": self.module,
            "full_content": self.full_content,
            "incremental_content": self.incremental_content,
            "has_incremental": self.has_incremental,
            "related_pages": self.related_pages,
            "prerequisite_pages": self.prerequisite_pages,
            "subsequent_pages": self.subsequent_pages,
            "raw_html": self.raw_html
        }


class MetadataFilter:
    """元数据过滤器工具类（符合ChromaDB where语法）"""

    @staticmethod
    def page_key(page_key: str) -> dict:
        """按页面key过滤"""
        return {"page_key": {"$eq": page_key}}

    @staticmethod
    def page_incremental(page_key: str) -> dict:
        """按页面key过滤增量"""
        return {
            "$and": [
                {"page_key": {"$eq": page_key}},
                {"is_incremental": {"$eq": True}}
            ]
        }

    @staticmethod
    def page_full(page_key: str) -> dict:
        """按页面key过滤全量内容"""
        return {
            "$and": [
                {"page_key": {"$eq": page_key}},
                {"doc_type": {"$eq": "page_full"}}
            ]
        }

    @staticmethod
    def module(module_name: str) -> dict:
        """按模块过滤"""
        return {
            "$and": [
                {"doc_type": {"$eq": "module_fused"}},
                {"module": {"$eq": module_name}}
            ]
        }

    @staticmethod
    def incremental_only() -> dict:
        """只要增量内容"""
        return {"is_incremental": {"$eq": True}}

    @staticmethod
    def doc_type(doc_type: str) -> dict:
        """按文档类型过滤"""
        return {"doc_type": {"$eq": doc_type}}

    @staticmethod
    def combine(*filters) -> dict:
        """组合多个过滤条件"""
        conditions = []
        for f in filters:
            if isinstance(f, dict):
                conditions.append(f)
        if len(conditions) == 1:
            return conditions[0]
        elif len(conditions) > 1:
            return {"$and": conditions}
        return {}


# ============ 蓝色字体识别 ============

def extract_blue_text_from_html(html_content: str) -> tuple:
    """
    从HTML中提取蓝色字体（增量）和非蓝色字体（全量）

    Args:
        html_content: HTML内容

    Returns:
        (full_content, incremental_content)
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        print("[蓝色字体识别] BeautifulSoup未安装，使用简单正则表达式")
        return extract_blue_text_simple(html_content)

    try:
        soup = BeautifulSoup(html_content, 'html.parser')

        blue_texts = []
        normal_texts = []

        # 遍历所有可能包含文本的元素
        for element in soup.find_all(['span', 'div', 'p', 'td', 'li', 'font', 'a']):
            style = element.get('style', '').lower().replace(' ', '')
            class_attr = element.get('class', [])
            color_attr = element.get('color', '').lower() if element.name == 'font' else ''

            # 判断是否为蓝色字体（多种判断方式）
            is_blue = (
                'color:blue' in style or
                'color:#0000ff' in style or
                'color:#00f' in style or
                'rgb(0,0,255)' in style or
                'color:rgb(0,0,255)' in style or
                any('blue' in str(c).lower() for c in class_attr) or
                color_attr == 'blue' or
                color_attr == '#0000ff' or
                color_attr == '#00f'
            )

            if is_blue:
                # 提取蓝色文本
                text = element.get_text(strip=True)
                if text and len(text) > 1:  # 过滤单字符
                    blue_texts.append(text)
            else:
                # 对于非蓝色元素，需要排除其蓝色子元素
                # 克隆元素以避免修改原始DOM
                from copy import deepcopy
                element_copy = deepcopy(element)

                # 移除蓝色的子元素
                for blue_elem in element_copy.find_all(style=lambda s: s and 'blue' in s.lower()):
                    blue_elem.decompose()
                for blue_elem in element_copy.find_all(class_=lambda c: c and any('blue' in str(cl).lower() for cl in c)):
                    blue_elem.decompose()

                text = element_copy.get_text(strip=True)
                if text and len(text) > 1:
                    normal_texts.append(text)

        # 合并文本
        full_content = "\n".join(normal_texts)
        incremental_content = "\n".join(blue_texts)

        print(f"[蓝色字体识别] 全量内容: {len(full_content)} 字符, 增量内容: {len(incremental_content)} 字符")

        return full_content, incremental_content

    except Exception as e:
        print(f"[蓝色字体识别] BeautifulSoup解析失败: {e}，使用简单正则表达式")
        return extract_blue_text_simple(html_content)


def extract_blue_text_simple(html_content: str) -> tuple:
    """
    简单版蓝色字体识别（使用正则表达式）

    当BeautifulSoup不可用时的降级方案
    """
    import re

    # 匹配蓝色字体的HTML标签
    # 匹配 style="...color:blue..." 或 style="...color:#0000ff..."
    blue_pattern = re.compile(
        r'<[^>]*(?:style\s*=\s*["\'][^"\']*?\bcolor\s*:\s*(?:blue|#0{0,2}00ff|#00f|rgb\(\s*0\s*,\s*0\s*,\s*255\s*\))[^"\']*?["\']|color\s*=\s*["\']?(?:blue|#0{0,2}00ff|#00f)["\']?)[^>]*>(.*?)</[^>]+>',
        re.IGNORECASE | re.DOTALL
    )

    blue_matches = blue_pattern.findall(html_content)
    incremental_content = "\n".join([m.strip() for m in blue_matches if m.strip()])

    # 移除蓝色标签后得到全量内容
    full_content = blue_pattern.sub('', html_content)

    # 清理HTML标签
    full_content = re.sub(r'<[^>]+>', '\n', full_content)
    full_content = re.sub(r'\n+', '\n', full_content).strip()

    return full_content, incremental_content


# ============ 格式化为Markdown ============

def format_page_to_markdown(page_data: PageData) -> dict:
    """
    将页面内容格式化为Markdown

    Args:
        page_data: PageData对象

    Returns:
        {"full_markdown": "...", "incremental_markdown": "..."}
    """
    # 格式化全量内容
    full_markdown = f"# {page_data.page_name}\n\n"
    full_markdown += page_data.full_content

    # 格式化增量内容
    incremental_markdown = ""
    if page_data.incremental_content:
        incremental_markdown = f"# {page_data.page_name} - 本次新增\n\n"
        incremental_markdown += page_data.incremental_content

    return {
        "page_key": page_data.page_key,
        "full_markdown": full_markdown,
        "incremental_markdown": incremental_markdown
    }


# ============ 原有知识库管理功能（保留兼容） ============

# 确保知识库目录存在
os.makedirs(KNOWLEDGE_BASE_PATH, exist_ok=True)


async def create_knowledge_base(
    name: str,
    axure_url: str,
    vision_provider: str = "doubao",
    username: str = None,
    password: str = None,
    use_ai_refine: bool = False,
    max_concurrent: int = 5
) -> Dict[str, Any]:
    """
    创建知识库（快速版）：直接 HTTP 获取 Axure 内容 -> 并发处理 -> 向量化存储

    不使用浏览器和截图，直接通过 HTTP 获取文本内容，速度快很多。

    Args:
        name: 知识库名称
        axure_url: Axure 原型首页链接
        vision_provider: 视觉模型提供商（此版本不使用）
        username: 登录用户名（此版本不使用）
        password: 登录密码（此版本不使用）
        use_ai_refine: 是否使用AI整理内容（默认关闭以加快速度）
        max_concurrent: 最大并发数

    Returns:
        Dict: 创建结果
    """
    import chromadb
    import asyncio
    import aiohttp
    from urllib.parse import urlparse, unquote, quote
    from bs4 import BeautifulSoup

    try:
        print(f"[知识库] 开始创建知识库: {name}")
        start_time = asyncio.get_event_loop().time()

        # 解析基础 URL
        parsed = urlparse(axure_url)
        base_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"

        pages_to_process = []
        all_documents = []

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        async with aiohttp.ClientSession(headers=headers) as session:
            # 1. 直接获取 sitemap.js 解析站点地图
            print("[知识库] 获取站点地图...")
            sitemap_js_url = base_url.replace('start.html', 'data/sitemap.js')

            try:
                async with session.get(sitemap_js_url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                    if resp.status == 200:
                        js_content = await resp.text()

                        # 解析 sitemap.js
                        match = re.search(r'\$axure\.loadSitemap\s*\(\s*(\{[\s\S]*\})\s*\)', js_content)
                        if match:
                            sitemap_json = match.group(1)
                            # 处理 JavaScript 对象转 JSON
                            sitemap_json = re.sub(r'(\w+)\s*:', r'"\1":', sitemap_json)
                            sitemap_json = sitemap_json.replace("'", '"')

                            try:
                                sitemap_data = json.loads(sitemap_json)

                                # 递归提取所有页面
                                def extract_pages(nodes):
                                    result = []
                                    for node in nodes:
                                        page_name = node.get('pageName') or node.get('name') or node.get('id', '')
                                        if page_name:
                                            result.append({
                                                'name': page_name,
                                                'url': f"{base_url}?p={page_name}",
                                                'page_id': page_name
                                            })
                                        if node.get('children'):
                                            result.extend(extract_pages(node['children']))
                                    return result

                                root_nodes = sitemap_data.get('rootNodes', sitemap_data.get('children', []))
                                pages_to_process = extract_pages(root_nodes)
                                print(f"[知识库] 从 sitemap.js 获取到 {len(pages_to_process)} 个页面")
                            except json.JSONDecodeError as e:
                                print(f"[知识库] sitemap JSON 解析失败: {e}")
                    else:
                        print(f"[知识库] sitemap.js 请求失败: {resp.status}")
            except Exception as e:
                print(f"[知识库] 获取站点地图失败: {e}")

            # 如果没有获取到页面，使用当前URL
            if not pages_to_process:
                current_page_name = unquote(parsed.query.split('p=')[-1]) if 'p=' in axure_url else "首页"
                pages_to_process.append({
                    'name': current_page_name,
                    'url': axure_url,
                    'page_id': current_page_name
                })
                print(f"[知识库] 使用当前页面: {current_page_name}")

            # 2. 并发获取每个页面的内容
            print(f"[知识库] 开始并发获取 {len(pages_to_process)} 个页面内容...")

            semaphore = asyncio.Semaphore(max_concurrent)

            async def fetch_page_content(page_info):
                """直接 HTTP 获取页面内容并解析"""
                async with semaphore:
                    try:
                        page_name = page_info['name']
                        # 构建页面 HTML URL
                        base_path = parsed.path.rsplit('/', 1)[0] if '/' in parsed.path else parsed.path
                        page_html_url = f"{parsed.scheme}://{parsed.netloc}{base_path}/{page_name}.html"

                        async with session.get(page_html_url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                            if resp.status != 200:
                                print(f"[知识库] ✗ {page_name}: HTTP {resp.status}")
                                return None

                            html_content = await resp.text()

                        # 使用 BeautifulSoup 解析
                        soup = BeautifulSoup(html_content, 'html.parser')

                        # 移除脚本和样式
                        for tag in soup(['script', 'style', 'noscript', 'link', 'meta']):
                            tag.decompose()

                        # 提取文本和表格
                        texts = []
                        tables_data = []

                        # 提取表格
                        for table in soup.find_all('table'):
                            rows = []
                            for tr in table.find_all('tr'):
                                cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                                if cells and any(c for c in cells):
                                    rows.append(cells)
                            if rows:
                                tables_data.append(rows)

                        # 提取文本元素
                        seen_texts = set()
                        for element in soup.find_all(['div', 'span', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'label', 'a']):
                            text = element.get_text(strip=True)
                            if text and len(text) > 1 and text not in seen_texts:
                                text_lower = text.lower()
                                if not any(kw in text_lower for kw in ['axure', 'prototype', 'preview', 'inspect', 'console']):
                                    seen_texts.add(text)
                                    texts.append(text)

                        # 构建 Markdown
                        md_parts = [f"# {page_name}\n"]

                        # 输出表格
                        for idx, table_rows in enumerate(tables_data):
                            if len(table_rows) >= 1:
                                max_cols = max(len(row) for row in table_rows)
                                header = table_rows[0]
                                header_padded = header + [''] * (max_cols - len(header))
                                md_parts.append("| " + " | ".join(header_padded) + " |")
                                md_parts.append("| " + " | ".join(['---'] * max_cols) + " |")
                                for row in table_rows[1:]:
                                    row_padded = row + [''] * (max_cols - len(row))
                                    row_cleaned = [cell.replace('\n', ' ').replace('|', '｜') for cell in row_padded]
                                    md_parts.append("| " + " | ".join(row_cleaned) + " |")
                                md_parts.append("")

                        # 输出文本
                        for text in texts[:200]:
                            md_parts.append(f"- {text}")

                        markdown_content = "\n".join(md_parts)

                        if markdown_content and len(markdown_content) > 50:
                            print(f"[知识库] ✓ {page_name}: {len(markdown_content)} 字符")
                            return {
                                'page_name': page_name,
                                'page_url': page_info['url'],
                                'page_id': page_info['page_id'],
                                'combined_content': markdown_content
                            }
                        else:
                            print(f"[知识库] ✗ {page_name}: 内容过少")
                            return None

                    except Exception as e:
                        print(f"[知识库] ✗ {page_info['name']}: {e}")
                        return None

            # 并发获取所有页面
            tasks = [fetch_page_content(page) for page in pages_to_process]
            results = await asyncio.gather(*tasks)

            # 过滤掉失败的结果
            all_documents = [doc for doc in results if doc is not None]
            print(f"[知识库] 成功获取 {len(all_documents)}/{len(pages_to_process)} 个页面")

        if not all_documents:
            return {
                "success": False,
                "error": "没有成功获取任何页面内容"
            }

        # 3. 可选的AI整理
        if use_ai_refine:
            print("[知识库] 开始AI整理内容...")
            for i, doc in enumerate(all_documents):
                try:
                    print(f"[知识库] AI整理 {i+1}/{len(all_documents)}: {doc['page_name']}")
                    refined = await refine_requirements_markdown_async(
                        doc['combined_content'],
                        provider="deepseek"
                    )
                    if refined and len(refined) > 100:
                        doc['combined_content'] = refined
                except Exception as e:
                    print(f"[知识库] AI整理失败: {e}")

        # 4. 分块处理（DashScope Embedding 限制单条最大8192字符）
        print("[知识库] 开始分块...")
        all_chunks = []
        for doc in all_documents:
            content = doc['combined_content']
            if len(content) > 7000:
                chunks = split_text_into_chunks(content, chunk_size=6000, chunk_overlap=200)
                for i, chunk in enumerate(chunks):
                    all_chunks.append({
                        'page_name': doc['page_name'],
                        'page_url': doc['page_url'],
                        'page_id': doc['page_id'],
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'content': chunk
                    })
            else:
                all_chunks.append({
                    'page_name': doc['page_name'],
                    'page_url': doc['page_url'],
                    'page_id': doc['page_id'],
                    'chunk_index': 0,
                    'total_chunks': 1,
                    'content': content
                })

        print(f"[知识库] 分块完成，共 {len(all_chunks)} 个文本块")

        # 5. 向量化
        print("[知识库] 开始向量化...")
        texts_to_embed = [chunk['content'] for chunk in all_chunks]
        embeddings = get_dashscope_embedding(texts_to_embed)

        # 6. 存储到 Chroma
        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)

        timestamp = get_current_datetime()
        name_hash = hashlib.md5(name.encode('utf-8')).hexdigest()[:8]
        collection_name = f"kb_{name_hash}_{timestamp}"

        collection = client.get_or_create_collection(
            name=collection_name,
            metadata={
                "kb_name": name,
                "display_name": name,
                "axure_url": axure_url,
                "created_at": timestamp,
                "type": "knowledge_base",
                "page_count": len(all_documents)
            }
        )

        ids = [f"chunk_{i}" for i in range(len(all_chunks))]
        metadatas = [{
            "page_name": chunk['page_name'],
            "page_url": chunk['page_url'],
            "page_id": chunk['page_id'],
            "chunk_index": chunk['chunk_index'],
            "total_chunks": chunk['total_chunks']
        } for chunk in all_chunks]
        documents = [chunk['content'] for chunk in all_chunks]

        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas
        )

        end_time = asyncio.get_event_loop().time()
        elapsed = end_time - start_time

        print(f"[知识库] 创建完成! 耗时: {elapsed:.1f}秒")

        return {
            "success": True,
            "collection_name": collection_name,
            "kb_name": name,
            "page_count": len(all_documents),
            "total_pages": len(all_documents),
            "total_chunks": len(all_chunks),
            "elapsed_seconds": round(elapsed, 1),
            "pages": [{"name": doc['page_name'], "url": doc['page_url']} for doc in all_documents],
            "message": f"知识库 '{name}' 创建成功，共处理 {len(all_documents)} 个页面，生成 {len(all_chunks)} 个文本块，耗时 {elapsed:.1f} 秒"
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": f"创建知识库失败: {str(e)}"
        }


def list_knowledge_bases() -> List[Dict[str, Any]]:
    """
    列出所有知识库

    Returns:
        List[Dict]: 知识库列表
    """
    try:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        collections = client.list_collections()

        knowledge_bases = []
        for col in collections:
            metadata = col.metadata or {}
            # 只返回知识库类型的 collection
            if metadata.get("type") == "knowledge_base":
                knowledge_bases.append({
                    "collection_name": col.name,
                    "kb_name": metadata.get("kb_name", col.name),
                    "axure_url": metadata.get("axure_url", ""),
                    "created_at": metadata.get("created_at", ""),
                    "page_count": metadata.get("page_count", col.count()),
                    "doc_count": col.count()
                })

        return knowledge_bases
    except Exception as e:
        print(f"[知识库] 列出知识库失败: {e}")
        return []


def delete_knowledge_base(collection_name: str) -> bool:
    """
    删除知识库

    Args:
        collection_name: collection 名称

    Returns:
        bool: 是否删除成功
    """
    try:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        client.delete_collection(name=collection_name)
        return True
    except Exception as e:
        print(f"[知识库] 删除知识库失败: {e}")
        return False


async def recall_from_knowledge_base(
    collection_name: str,
    query_text: str,
    top_k: int = 5
) -> Dict[str, Any]:
    """
    从知识库召回相关内容

    Args:
        collection_name: collection 名称
        query_text: 查询文本
        top_k: 返回数量

    Returns:
        Dict: 召回结果
    """
    try:
        result = query_chroma_vector_db(
            collection_name=collection_name,
            query_text=query_text,
            top_k=top_k
        )
        return result
    except Exception as e:
        return {
            "success": False,
            "error": f"召回失败: {str(e)}"
        }


async def generate_test_points(
    context: str,
    requirement: str,
    provider: str = "azure"
) -> Dict[str, Any]:
    """
    根据召回内容生成测试点

    Args:
        context: 召回的上下文内容
        requirement: 用户需求描述
        provider: LLM 提供商

    Returns:
        Dict: 测试点列表
    """
    from llms import call_llm_api

    prompt = f"""你是一名资深测试工程师。请根据以下需求信息，提取关键的测试点。

## 用户需求
{requirement}

## 相关功能文档
{context}

## 输出要求
请列出所有需要测试的测试点，每个测试点应包含：
1. 测试点名称（简洁明了）
2. 测试目的（验证什么功能/场景）
3. 测试类型（功能测试/边界测试/异常测试/性能测试等）

请直接输出 JSON 格式：
{{
  "test_points": [
    {{
      "name": "测试点名称",
      "purpose": "测试目的",
      "type": "测试类型"
    }}
  ]
}}
"""

    try:
        response = await call_llm_api(prompt, provider=provider)
        content = response.get("choices", [{}])[0].get("message", {}).get("content", "")

        # 尝试解析 JSON
        import json
        # 提取 JSON 部分
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            test_points = json.loads(json_match.group())
            return {
                "success": True,
                "test_points": test_points.get("test_points", []),
                "raw_content": content
            }
        else:
            return {
                "success": True,
                "test_points": [],
                "raw_content": content
            }
    except Exception as e:
        return {
            "success": False,
            "error": f"生成测试点失败: {str(e)}"
        }


async def generate_test_cases_from_points(
    test_points: List[Dict],
    context: str,
    test_module: str,
    provider: str = "azure"
) -> Dict[str, Any]:
    """
    根据测试点生成详细测试用例

    Args:
        test_points: 测试点列表
        context: 上下文内容
        test_module: 模块名称
        provider: LLM 提供商

    Returns:
        Dict: 测试用例
    """
    from llms import call_llm_api

    # 格式化测试点
    points_text = "\n".join([
        f"{i+1}. {p['name']} - {p['purpose']} ({p['type']})"
        for i, p in enumerate(test_points)
    ])

    prompt = f"""你是一名资深测试工程师。请根据以下测试点，生成详细的测试用例。

## 模块名称
{test_module}

## 测试点列表
{points_text}

## 相关功能文档
{context}

## 输出要求
为每个测试点生成 1-2 个测试用例，每个测试用例包含：
- case_id: 用例ID（格式 TC-001, TC-002...）
- test_module: 模块名称
- title: 用例标题
- precondition: 前置条件
- steps: 测试步骤（每步一行，用换行符分隔）
- expected_result: 预期结果（每项一行，用换行符分隔）
- priority: 优先级（高/中/低）

请直接输出 JSON 格式：
{{
  "test_cases": [
    {{
      "case_id": "TC-001",
      "test_module": "{test_module}",
      "title": "用例标题",
      "precondition": "前置条件",
      "steps": "1. 步骤1\\n2. 步骤2",
      "expected_result": "1. 预期1\\n2. 预期2",
      "priority": "高"
    }}
  ]
}}
"""

    try:
        response = await call_llm_api(prompt, provider=provider)
        content = response.get("choices", [{}])[0].get("message", {}).get("content", "")

        # 修复 JSON 控制字符
        fixed_content = _fix_json_control_chars(content)

        # 提取 JSON 部分
        json_match = re.search(r'\{[\s\S]*\}', fixed_content)
        if json_match:
            test_cases = json.loads(json_match.group())
            return {
                "success": True,
                "test_cases": test_cases,
                "raw_content": content
            }
        else:
            return {
                "success": True,
                "test_cases": {"test_cases": []},
                "raw_content": content
            }
    except Exception as e:
        return {
            "success": False,
            "error": f"生成测试用例失败: {str(e)}"
        }
# -*- coding: utf-8 -*-
"""
增强版知识库管理功能 - Phase 2-5
包含：
- Phase 2: 增强版Axure解析
- Phase 3: 分层向量化存储
- Phase 4: 大模型关联分析
- Phase 5: 智能召回系统
"""

import os
import json
import asyncio
import time
from datetime import datetime
from typing import Dict, Any, List, Optional
from urllib.parse import urlparse, unquote, parse_qs

import aiohttp
from bs4 import BeautifulSoup

# 导入基础工具
from utils import (
    PageData, MetadataFilter,
    extract_blue_text_from_html,
    format_page_to_markdown,
    CHROMA_DB_PATH,
    call_llm_api,
    KNOWLEDGE_BASE_PATH
)

# ============ Phase 5: 智能召回系统 ============
def add_title_suffix_to_documents(documents: list, suffix: str) -> list:
    """
    给文档列表中的第一个标题添加后缀

    Args:
        documents: 文档列表
        suffix: 要添加的后缀（如 " - 当前页面全量"）

    Returns:
        修改后的文档列表
    """
    if not documents or not documents[0]:
        return documents

    result = []
    for doc in documents:
        if doc:
            lines = doc.split('\n')
            modified = False
            for i, line in enumerate(lines):
                stripped = line.strip()
                # 找到第一个 # 开头的标题
                if stripped.startswith('#') and not modified:
                    # 在标题后添加后缀（去掉换行符）
                    lines[i] = line.rstrip() + ' ' + suffix
                    modified = True
                    break
            result.append('\n'.join(lines))
        else:
            result.append(doc)

    return result


def assemble_incremental_with_context(
    incremental: dict,
    page_full: dict,
    context: dict,
    query: str
) -> str:
    """
    组装增量 + 全量页面 + 模块的完整prompt

    Args:
        incremental: 增量召回结果（蓝色内容）
        page_full: 页面全量召回结果
        context: 模块融合召回结果
        query: 用户查询

    Returns:
        组装后的完整内容
    """
    parts = []

    # 1. 添加模块背景（最宏观）
    if context and context.get('documents'):
        parts.append(f"""#  模块功能背景（整体业务逻辑）

请首先理解以下模块的整体业务逻辑和基础规则：


{context['documents'][0]}\n


""")
        print(f"[组装-DEBUG] 添加模块背景: {len(context['documents'][0])} 字符")
    else:
        print(f"[组装-DEBUG] 跳过模块背景: context={context}, documents={context.get('documents') if context else None}")

    # 2. 添加页面全量内容
    if page_full and page_full.get('documents'):
        page_name = page_full.get('metadatas', [{}])[0].get('page_name', '当前页面')
        parts.append(f"""

以下是该页面的完整功能说明：


{page_full['documents'][0]}\n


""")
        print(f"[组装-DEBUG] 添加页面全量: {len(page_full['documents'][0])} 字符")
    else:
        print(f"[组装-DEBUG] 跳过页面全量: page_full={page_full}, documents={page_full.get('documents') if page_full else None}")

    # 3. 添加增量需求（蓝色内容）
    if incremental.get('documents'):
        parts.append(f"""

以下是需要重点测试的本次新增或变更的功能点：


{incremental['documents'][0]}\n


""")
        print(f"[组装-DEBUG] 添加增量需求: {len(incremental['documents'][0])} 字符")
    else:
        print(f"[组装-DEBUG] 跳过增量需求: documents={incremental.get('documents')}")

    result = "\n".join(parts)
    print(f"[组装-DEBUG] 最终组装: {len(parts)} 个部分, 总长度 {len(result)} 字符")
    return result


# ============ AI结构化召回内容 ============

async def structure_recall_content_with_ai(combined_content: str, query: str, provider: str = "deepseek") -> dict:
    """
    使用AI对召回的内容进行结构化处理

    Args:
        combined_content: 组装后的内容（已格式化的 Markdown）
        query: 用户查询
        provider: LLM提供商

    Returns:
        结构化后的结果 {
            "structured": "AI结构化后的内容",
            "raw": "原始 Markdown 内容"
        }
    """
    if not combined_content:
        return {
            "structured": "",
            "raw": ""
        }

    raw_content = combined_content

    try:
        print(f"[AI结构化] 开始处理内容, 总长度={len(raw_content)}")

        # 直接调用 AI 结构化
        structured_content = await refine_requirements_markdown_async(raw_content, provider=provider)

        print(f"[AI结构化] 完成, structured 长度={len(structured_content)}")

        return {
            "structured": structured_content,
            "raw": raw_content
        }

    except Exception as e:
        print(f"[AI结构化] 结构化失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            "structured": raw_content,  # 失败时返回原始内容
            "raw": raw_content
        }


# ============ Phase 5: 智能召回系统 ============

def format_recall_results(query_results: dict) -> dict:
    """
    格式化召回结果

    Args:
        query_results: ChromaDB查询结果

    Returns:
        格式化后的结果
    """
    if not query_results or not query_results.get('ids') or not query_results['ids'][0]:
        return {"documents": [], "metadatas": []}

    return {
        "documents": query_results.get('documents', [[]])[0],
        "metadatas": query_results.get('metadatas', [[]])[0],
        "distances": query_results.get('distances', [[]])[0]
    }


async def smart_recall_from_knowledge_base(
    collection_name: str,
    query: str,
    page_key: Optional[str] = None,
    recall_strategy: str = "incremental_with_context",
    top_k: int = 5,
    use_ai_structure: bool = False,
    ai_provider: str = "deepseek"
) -> dict:
    """
    智能召回：根据查询内容自动选择最佳召回策略

    Args:
        collection_name: collection名称
        query: 查询文本
        page_key: 页面key（可选）
        recall_strategy: 召回策略 (auto | incremental_with_context | page_level | module_level)
        top_k: 召回数量
        use_ai_structure: 是否使用AI结构化召回内容
        ai_provider: AI服务提供商

    Returns:
        召回结果，包含原始内容和AI结构化内容（如果启用）
    """
    try:
        import chromadb

        client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        collection = client.get_collection(collection_name)

        print(f"[智能召回] 召回策略: {recall_strategy}, 查询: {query[:50]}...")

        # 🔥 关键：使用 DashScope embedding 对查询文本进行向量化
        query_embedding = get_dashscope_embedding([query])[0]
        print(f"[智能召回] Query embedding 维度: {len(query_embedding)}")


        results = {"strategy": recall_strategy}

        if recall_strategy == "incremental_with_context":
            # ========== 策略1: 蓝色 + 全量 + 模块三合一 ==========

            # 1.1 召回增量内容（蓝色）
            where_filter = MetadataFilter.incremental_only()
            if page_key:
                where_filter = MetadataFilter.page_incremental(page_key)

            incremental_results = collection.query(
                query_embeddings=[query_embedding],
                where=where_filter,
                n_results=top_k
            )
            # 格式化
            incremental_formatted = format_recall_results(incremental_results)

            # 1.2 从增量结果中提取 page_key，并召回页面全量内容
            page_full_formatted = None
            actual_page_key = page_key  # 先使用传入的 page_key

            # 如果增量召回有结果，从 metadata 中提取 page_key
            if incremental_formatted.get('metadatas') and not actual_page_key:
                page_metadata = incremental_formatted['metadatas'][0]
                actual_page_key = page_metadata.get('page_key', '')
                print(f"[智能召回-DEBUG] 从增量metadata提取page_key: '{actual_page_key}'")

            if actual_page_key:
                print(f"[智能召回-DEBUG] 准备查询页面全量: page_key='{actual_page_key}'")
                page_full_results = collection.query(
                    query_embeddings=[query_embedding],
                    where=MetadataFilter.page_full(actual_page_key),
                    n_results=top_k
                )
                # 格式化
                page_full_formatted = format_recall_results(page_full_results)
                # 给页面全量的第一个标题添加后缀
                # if page_full_formatted.get('documents'):
                #     page_full_formatted['documents'] = add_title_suffix_to_documents(
                #         page_full_formatted['documents'],
                #         "- 当前页面全量"
                #     )
                print(f"[智能召回-DEBUG] 页面全量查询结果documents数量: {len(page_full_formatted.get('documents', []))}")

                # 调试：打印页面全量内容的前500字符，检查是否有表格
                if page_full_formatted.get('documents'):
                    sample = page_full_formatted['documents'][0][:500]
                    print(f"[智能召回-DEBUG] 页面全量内容样本:\n{sample}")
                    has_table = '|' in sample and '---' in sample
                    print(f"[智能召回-DEBUG] 检测到表格格式: {has_table}")
            else:
                print(f"[智能召回-DEBUG] 未获取到page_key，跳过页面全量查询")

            # 1.3 获取模块上下文
            print(f"[智能召回-DEBUG] 开始获取模块上下文...")
            print(f"[智能召回-DEBUG] incremental metadatas数量: {len(incremental_formatted.get('metadatas', []))}")

            if incremental_formatted.get('metadatas'):
                page_metadata = incremental_formatted['metadatas'][0]
                module_name = page_metadata.get('module', '')

                print(f"[智能召回-DEBUG] 从增量metadata获取模块名: '{module_name}'")
                print(f"[智能召回-DEBUG] 增量metadata: {page_metadata}")

                context_formatted = None
                if module_name:
                    print(f"[智能召回-DEBUG] 准备查询模块: '{module_name}'")
                    print(f"[智能召回-DEBUG] 过滤条件: {MetadataFilter.module(module_name)}")
                    # 对模块名也进行向量化
                    module_embedding = get_dashscope_embedding([module_name])[0]
                    context_results = collection.query(
                        query_embeddings=[module_embedding],
                        where=MetadataFilter.module(module_name),
                        n_results=1
                    )

                    print(f"[智能召回-DEBUG] 模块查询结果documents数量: {len(context_results.get('documents', [[]])[0]) if context_results.get('documents') else 0}")
                    # 格式化
                    context_formatted = format_recall_results(context_results)
                    # 给模块关联需求的第一个标题添加后缀
                    if context_formatted.get('documents'):
                        context_formatted['documents'] = add_title_suffix_to_documents(
                            context_formatted['documents'],
                            "- 关联需求"
                        )
                    print(f"[智能召回-DEBUG] 格式化后documents数量: {len(context_formatted.get('documents', []))}")

                # 组装结果：蓝色 + 全量 + 模块
                results['incremental'] = incremental_formatted
                results['page_full'] = page_full_formatted
                results['context'] = context_formatted

                # 打印调试信息
                inc_docs = len(incremental_formatted.get('documents', [])) if incremental_formatted else 0
                pf_docs = len(page_full_formatted.get('documents', [])) if page_full_formatted else 0
                ctx_docs = len(context_formatted.get('documents', [])) if context_formatted else 0
                print(f"[智能召回-DEBUG] 组装前: incremental={inc_docs} docs, page_full={pf_docs} docs, context={ctx_docs} docs")

                results['combined'] = assemble_incremental_with_context(
                    incremental=incremental_formatted,
                    page_full=page_full_formatted,
                    context=context_formatted,
                    query=query
                )

                print(f"[智能召回-DEBUG] combined 内容长度: {len(results['combined'])} 字符")
                print(f"[智能召回-DEBUG] combined 前200字符: {results['combined'][:200]}")

                results['total_chunks'] = len(incremental_formatted.get('documents', []))
            else:
                # 没有增量，降级为普通召回
                print(f"[智能召回] 未找到增量内容，降级为 page_level")
                return await smart_recall_from_knowledge_base(
                    collection_name=collection_name,
                    query=query,
                    page_key=page_key,
                    recall_strategy="page_level",
                    top_k=top_k
                )

        # 🆕 AI结构化处理（如果启用）
        if use_ai_structure:
            # 优先处理 combined（增量+全量+模块），其次处理 data（普通召回）
            content_to_structure = None

            if results.get('combined'):
                # incremental_with_context 策略返回的是 combined
                content_to_structure = results['combined']
                print(f"[智能召回] 使用 combined 内容进行AI结构化...")
            elif results.get('data') and results['data'].get('documents'):
                # 其他策略返回的是 data.documents，先合并
                documents = results['data']['documents']
                content_to_structure = "\n\n---\n\n".join(documents)
                print(f"[智能召回] 使用 data 内容进行AI结构化...")

            if content_to_structure:
                print(f"[智能召回] 正在进行AI结构化处理...")
                print(f"[智能召回-DEBUG] 待处理内容长度: {len(content_to_structure)} 字符")
                structured = await structure_recall_content_with_ai(content_to_structure, query, ai_provider)
                results['structured'] = structured['structured']
                results['raw'] = structured['raw']  # 格式化后的 Markdown
                print(f"[智能召回] AI结构化完成")
                print(f"[智能召回-DEBUG] raw 内容长度: {len(structured['raw'])} 字符")

        return results

    except Exception as e:
        print(f"[智能召回] 召回失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            "strategy": recall_strategy,
            "error": str(e)
        }
