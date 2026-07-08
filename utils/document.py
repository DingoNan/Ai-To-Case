"""文档解析模块

包含 Markdown、PDF、Word 文档的解析功能。
"""
import os
import re
import asyncio
from .ocr import ocr_image_async


async def parse_markdown_with_images(content: str, base_path: str = None, vision_provider: str = "aliyun") -> str:
    """
    解析 Markdown 文件中的图片引用，调用视觉大模型识别图片内容，
    并将识别结果还原到原文本位置，标注为【图片解析】

    Args:
        content: Markdown 文件的文本内容
        base_path: MD 文件所在的基础路径（用于解析相对路径图片）
        vision_provider: 视觉模型提供商，"aliyun" 或 "doubao"

    Returns:
        str: 处理后的文本，图片引用被替换为识别结果
    """
    import aiohttp
    from urllib.parse import urlparse, urljoin

    # 匹配 Markdown 图片语法: ![alt](path/url) 或 ![alt](path/url "title")
    image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"[^"]*")?\)')

    # 找到所有图片引用
    matches = list(image_pattern.finditer(content))

    if not matches:
        return content

    print(f"[MD图片解析] 找到 {len(matches)} 个图片引用")

    async def fetch_and_ocr_image(match) -> tuple:
        """获取图片并进行 OCR 识别"""
        alt_text = match.group(1)
        image_path = match.group(2)
        original_text = match.group(0)

        try:
            image_bytes = None

            # 判断是 URL 还是本地路径
            parsed = urlparse(image_path)
            if parsed.scheme in ('http', 'https'):
                # 网络图片，使用 aiohttp 下载
                async with aiohttp.ClientSession() as session:
                    async with session.get(image_path, timeout=aiohttp.ClientTimeout(total=30)) as response:
                        if response.status == 200:
                            image_bytes = await response.read()
                        else:
                            print(f"[MD图片解析] 下载图片失败: {image_path}, 状态码: {response.status}")
                            return match.start(), original_text, None
            else:
                # 本地图片路径
                local_path = image_path

                # 如果是相对路径，结合 base_path
                if base_path and not os.path.isabs(image_path):
                    local_path = os.path.join(base_path, image_path)

                # 处理 Windows 路径
                local_path = local_path.replace('/', os.sep).replace('\\', os.sep)

                if os.path.exists(local_path):
                    with open(local_path, 'rb') as f:
                        image_bytes = f.read()
                else:
                    print(f"[MD图片解析] 本地图片不存在: {local_path}")
                    return match.start(), original_text, None

            if image_bytes:
                # 调用 OCR 识别，忽略 token 统计（文档解析场景无需记录）
                ocr_text, _ = await ocr_image_async(image_bytes, vision_provider=vision_provider)

                if ocr_text and ocr_text.strip():
                    # 构建替换文本，保留原图片引用并添加解析结果
                    alt_desc = f"（{alt_text}）" if alt_text else ""
                    replacement = f"\n\n【图片解析开始】{alt_desc}\n{ocr_text}\n【图片解析结束】\n\n"
                    return match.start(), original_text, replacement
                else:
                    print(f"[MD图片解析] 图片 OCR 无结果: {image_path}")
                    return match.start(), original_text, None

        except asyncio.TimeoutError:
            print(f"[MD图片解析] 下载图片超时: {image_path}")
        except Exception as e:
            print(f"[MD图片解析] 处理图片失败: {image_path}, 错误: {str(e)}")

        return match.start(), original_text, None

    # 并发处理所有图片
    tasks = [fetch_and_ocr_image(m) for m in matches]
    results = await asyncio.gather(*tasks)

    # 按位置倒序排列，从后往前替换（避免位置偏移）
    results_sorted = sorted(results, key=lambda x: x[0], reverse=True)

    # 执行替换
    result_content = content
    for start_pos, original_text, replacement in results_sorted:
        if replacement:
            result_content = result_content.replace(original_text, replacement, 1)

    success_count = sum(1 for _, _, r in results if r is not None)
    print(f"[MD图片解析] 成功解析 {success_count}/{len(matches)} 个图片")

    return result_content


async def parse_pdf_to_text(file_bytes: bytes, enable_ocr: bool = True) -> str:
    """
    解析 PDF 文件，提取文本内容和图片（OCR识别）
    图片 OCR 结果会被插入到对应页面的位置，标注为【图片解析】

    Args:
        file_bytes: PDF 文件的字节内容
        enable_ocr: 是否启用图片 OCR 识别

    Returns:
        str: 提取的文本内容（包含 OCR 识别的图片内容，位于对应页面）
    """
    try:
        import fitz  # PyMuPDF

        # 从字节创建 PDF 文档
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # 按页面组织内容：{page_num: {"text": str, "images": [(img_index, image_bytes)]}}
        page_contents = {}

        for page_num, page in enumerate(doc):
            page_contents[page_num] = {
                "text": "",
                "images": []
            }

            # 提取文本内容
            page_text = page.get_text()
            if page_text.strip():
                page_contents[page_num]["text"] = page_text

            # 提取图片
            if enable_ocr:
                page_images = page.get_images(full=True)
                for img_index, img in enumerate(page_images):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        page_contents[page_num]["images"].append((img_index, image_bytes))
                    except Exception as e:
                        print(f"提取图片失败: {e}")

        doc.close()

        # OCR 识别所有图片（并发处理）
        if enable_ocr:
            # 收集所有需要 OCR 的图片
            ocr_tasks = []
            task_info = []  # 记录每个任务对应的页面和图片索引

            for page_num, content in page_contents.items():
                for img_index, image_bytes in content["images"]:
                    task_info.append((page_num, img_index))
                    ocr_tasks.append(ocr_image_async(image_bytes, vision_provider="aliyun"))

            # 并发执行 OCR
            if ocr_tasks:
                ocr_results = await asyncio.gather(*ocr_tasks, return_exceptions=True)

                # 将 OCR 结果（取文本部分，忽略 token 统计）按页面组织
                page_ocr_results = {}
                for i, result in enumerate(ocr_results):
                    page_num, img_index = task_info[i]
                    if page_num not in page_ocr_results:
                        page_ocr_results[page_num] = []

                    if isinstance(result, Exception):
                        print(f"OCR 识别失败（第 {page_num + 1} 页，图片 {img_index + 1}）: {result}")
                    else:
                        ocr_text = result[0] if result else ""
                        if ocr_text and ocr_text.strip():
                            page_ocr_results[page_num].append((img_index, ocr_text))

                # 将 OCR 结果存入对应页面
                for page_num, results in page_ocr_results.items():
                    page_contents[page_num]["ocr_results"] = results

        # 组装最终文本，图片内容紧跟在对应页面文本之后
        text_parts = []
        for page_num in sorted(page_contents.keys()):
            content = page_contents[page_num]

            # 添加页面标题和文本
            page_text = content["text"]
            if page_text.strip():
                text_parts.append(f"### 第 {page_num + 1} 页\n{page_text}")

            # 在页面文本之后添加该页面的图片 OCR 结果
            ocr_results = content.get("ocr_results", [])
            for img_index, ocr_text in sorted(ocr_results, key=lambda x: x[0]):
                text_parts.append(f"\n【图片解析开始】（第 {page_num + 1} 页，图片 {img_index + 1}）\n{ocr_text}\n【图片解析结束】\n")

        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 解析失败: {str(e)}")


async def parse_word_to_text(file_bytes: bytes, enable_ocr: bool = True) -> str:
    """
    解析 Word 文档（.docx），提取文本内容、表格和图片（OCR识别）
    图片 OCR 结果会被插入到文档中图片的原始位置，标注为【图片解析】

    Args:
        file_bytes: Word 文件的字节内容
        enable_ocr: 是否启用图片 OCR 识别

    Returns:
        str: 提取的文本内容（包含表格和 OCR 识别的图片内容，位于原始位置）
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from io import BytesIO

        # 辅助函数：正确提取表格行数据，处理合并单元格
        def extract_row_data(row):
            """提取行数据，正确处理合并单元格（避免重复内容）"""
            row_data = []
            processed_cells = set()  # 记录已处理的单元格索引

            for col_idx, cell in enumerate(row.cells):
                # 检查是否已经处理过这个单元格（合并单元格的情况）
                if col_idx in processed_cells:
                    row_data.append("")  # 合并的单元格位置添加空字符串
                    continue

                # 获取单元格属性
                tc_pr = cell._element.tcPr
                if tc_pr is None:
                    # 没有特殊属性，直接添加文本
                    row_data.append(cell.text.strip())
                    continue

                # 检查水平合并 (gridSpan)
                grid_span_element = tc_pr.gridSpan
                if grid_span_element is not None:
                    span_count = int(grid_span_element.get('val', 1))
                    if span_count > 1:
                        # 这是水平合并的起始单元格
                        row_data.append(cell.text.strip())
                        # 标记被合并的列位置
                        for i in range(1, span_count):
                            processed_cells.add(col_idx + i)
                        continue

                # 检查垂直合并 (vMerge)
                v_merge_element = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if v_merge_element is not None:
                    v_merge_val = v_merge_element.get('val', 'continue')
                    if v_merge_val == 'continue':
                        # 这是垂直合并的延续部分
                        row_data.append("")
                        continue
                    else:
                        # 这是垂直合并的起始单元格
                        row_data.append(cell.text.strip())
                        continue

                # 检查水平合并延续 (hMerge)
                h_merge_element = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hMerge')
                if h_merge_element is not None:
                    # 这是水平合并的延续部分
                    row_data.append("")
                    continue

                # 普通单元格
                row_data.append(cell.text.strip())

            return row_data

        # 从字节创建文档
        doc = Document(BytesIO(file_bytes))

        # 1. 首先提取所有图片并建立映射
        image_map = {}  # {rId: image_bytes}
        if enable_ocr:
            for rel_id, rel in doc.part.rels.items():
                if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image":
                    try:
                        image_map[rel_id] = rel.target_part.blob
                        print(f"[Word解析] 发现图片: {rel_id}")
                    except Exception as e:
                        print(f"[Word解析] 提取图片失败 (rId={rel_id}): {e}")

        print(f"[Word解析] 共发现 {len(image_map)} 个图片")

        # 2. 按文档顺序收集内容
        content_items = []  # [(position, type, content)]
        position = 0
        image_tasks = []  # OCR 任务列表
        processed_images = set()  # 已处理的图片 rId

        # 定义命名空间
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        }

        # 遍历文档的 body 元素
        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':  # 段落
                # 提取段落文本
                para_text = ''
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text
                        if para.style and para.style.name:
                            style_name = para.style.name.lower()
                            if 'heading 1' in style_name or 'title' in style_name:
                                para_text = f"# {para.text}"
                            elif 'heading 2' in style_name:
                                para_text = f"## {para.text}"
                            elif 'heading 3' in style_name:
                                para_text = f"### {para.text}"
                        break

                if para_text.strip():
                    content_items.append((position, 'text', para_text))
                    position += 1

                # 查找段落中的图片
                if enable_ocr:
                    blips = element.findall('.//a:blip', namespaces)
                    if not blips:
                        # 备用方式：直接搜索包含 embed 属性的元素
                        for child in element.iter():
                            if child.tag.endswith('}blip'):
                                blips.append(child)

                    for blip in blips:
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if not embed_id:
                            embed_id = blip.get(qn('r:embed'))

                        if embed_id and embed_id in image_map and embed_id not in processed_images:
                            img_bytes = image_map[embed_id]
                            processed_images.add(embed_id)
                            content_items.append((position, 'image_placeholder', len(image_tasks)))
                            image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                            position += 1
                            print(f"[Word解析] 段落中发现图片引用: {embed_id}")

            elif tag == 'tbl':  # 表格
                for table in doc.tables:
                    if table._tbl == element:
                        rows_data = []
                        for row in table.rows:
                            row_data = extract_row_data(row)
                            rows_data.append(row_data)

                        if rows_data:
                            header = rows_data[0]
                            table_text = "| " + " | ".join(header) + " |\n"
                            table_text += "| " + " | ".join(["---"] * len(header)) + " |\n"
                            for row in rows_data[1:]:
                                table_text += "| " + " | ".join(row) + " |\n"
                            content_items.append((position, 'table', table_text))
                            position += 1
                        break

        # 3. 处理未被引用但存在的图片
        if enable_ocr:
            unprocessed_images = set(image_map.keys()) - processed_images
            if unprocessed_images:
                print(f"[Word解析] 发现 {len(unprocessed_images)} 个未在段落中引用的图片，追加处理")
                for rel_id in unprocessed_images:
                    img_bytes = image_map[rel_id]
                    content_items.append((position, 'image_placeholder', len(image_tasks)))
                    image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                    position += 1

        # 4. 如果没有提取到任何内容，使用简单备用方式
        if not content_items:
            print("[Word解析] 主要方式未提取到内容，使用备用方式")
            for para in doc.paragraphs:
                if para.text.strip():
                    para_text = para.text
                    if para.style and para.style.name:
                        style_name = para.style.name.lower()
                        if 'heading 1' in style_name or 'title' in style_name:
                            para_text = f"# {para.text}"
                        elif 'heading 2' in style_name:
                            para_text = f"## {para.text}"
                        elif 'heading 3' in style_name:
                            para_text = f"### {para.text}"
                    content_items.append((position, 'text', para_text))
                    position += 1

            for table_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    row_data = extract_row_data(row)
                    rows_data.append(row_data)
                if rows_data:
                    header = rows_data[0]
                    table_text = f"\n### 表格 {table_idx + 1}\n"
                    table_text += "| " + " | ".join(header) + " |\n"
                    table_text += "| " + " | ".join(["---"] * len(header)) + " |\n"
                    for row in rows_data[1:]:
                        table_text += "| " + " | ".join(row) + " |\n"
                    content_items.append((position, 'table', table_text))
                    position += 1

            # 所有图片追加到末尾
            if enable_ocr and image_map:
                for rel_id, img_bytes in image_map.items():
                    content_items.append((position, 'image_placeholder', len(image_tasks)))
                    image_tasks.append(ocr_image_async(img_bytes, vision_provider="aliyun"))
                    position += 1

        # 5. 并发执行 OCR
        ocr_results = {}
        if image_tasks:
            print(f"[Word解析] 开始 OCR 识别 {len(image_tasks)} 个图片...")
            results = await asyncio.gather(*image_tasks, return_exceptions=True)
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    print(f"[Word解析] 图片 OCR 失败: {result}")
                    ocr_results[i] = None
                else:
                    ocr_text = result[0] if result else ""
                    if ocr_text and ocr_text.strip():
                        ocr_results[i] = ocr_text
                        print(f"[Word解析] 图片 {i+1} OCR 成功")
                    else:
                        ocr_results[i] = None

        # 6. 组装最终内容
        text_parts = []
        image_counter = 0
        for pos, item_type, content in sorted(content_items, key=lambda x: x[0]):
            if item_type == 'text':
                text_parts.append(content)
            elif item_type == 'table':
                text_parts.append(content)
            elif item_type == 'image_placeholder':
                task_idx = content
                image_counter += 1
                ocr_text = ocr_results.get(task_idx)
                if ocr_text:
                    text_parts.append(f"\n【图片解析开始】（图片 {image_counter}）\n{ocr_text}\n【图片解析结束】\n")

        print(f"[Word解析] 完成，共处理 {image_counter} 个图片")
        return "\n\n".join(text_parts)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise Exception(f"Word 文档解析失败: {str(e)}")
